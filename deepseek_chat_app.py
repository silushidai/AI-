# -*- coding: utf-8 -*-
"""
DeepSeek 对话客户端 - 带思维链流程图展示
支持两种模式：1) DeepSeek API（需联网） 2) Ollama 本地模型（可选用本机已部署的模型）
融合 AI 外置记忆循环结构、单思路结点动态更新（参考 ai编程.pdf 思路论）
"""
from __future__ import unicode_literals

import sys
import io
import os
from datetime import datetime
# Windows 下强制使用 UTF-8，避免中文等字符导致 ascii 编码错误
if sys.platform == 'win32':
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    except (AttributeError, OSError):
        pass

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, font as tkfont, filedialog
import threading
import re
import json
import random
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError

try:
    import requests
except ImportError:
    requests = None

import sqlite3

try:
    import pymysql
except ImportError:
    pymysql = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

import base64

# 支持上传的文件扩展名（文本 + 图片）
_UPLOAD_EXT_TEXT = ('.txt', '.docx', '.doc')
_UPLOAD_EXT_IMAGE = ('.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif')
_UPLOAD_ALLOWED_EXT = _UPLOAD_EXT_TEXT + _UPLOAD_EXT_IMAGE


def _normalize_file_dialog_paths(paths):
    """将 askopenfilenames 的返回值统一转为路径列表。兼容返回 str 或 tuple 的情况。"""
    if not paths:
        return []
    if isinstance(paths, str):
        return [paths]
    return [p for p in paths if isinstance(p, str) and p.strip()]

_IMAGE_MIME = {
    '.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
    '.webp': 'image/webp', '.bmp': 'image/bmp', '.gif': 'image/gif',
}


def _read_uploaded_file(path):
    """
    读取上传文件。支持 .txt、.docx、.png、.jpg 等。
    返回 (success, result, error_msg)。
    文本文件：result 为字符串。
    图片文件：result 为 dict {'type':'image', 'base64':str, 'mime':str}。
    """
    path = os.path.abspath(path)
    if not os.path.isfile(path):
        return False, None, '文件不存在'
    ext = os.path.splitext(path)[1].lower()
    # 文本
    if ext == '.txt':
        for enc in ('utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1'):
            try:
                with open(path, 'r', encoding=enc) as f:
                    return True, f.read(), ''
            except (UnicodeDecodeError, OSError):
                continue
        return False, None, '无法识别文本编码'
    if ext == '.docx':
        if DocxDocument is None:
            return False, None, '请先安装: pip install python-docx'
        try:
            doc = DocxDocument(path)
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            return True, '\n'.join(parts), ''
        except Exception as e:
            return False, None, str(e)
    if ext == '.doc':
        return False, None, '请将 .doc 文件另存为 .docx 后上传'
    # 图片：转为 base64
    if ext in _UPLOAD_EXT_IMAGE:
        try:
            with open(path, 'rb') as f:
                b64 = base64.b64encode(f.read()).decode('utf-8')
            mime = _IMAGE_MIME.get(ext, 'image/png')
            return True, {'type': 'image', 'base64': b64, 'mime': mime}, ''
        except Exception as e:
            return False, None, str(e)
    return False, None, '不支持的格式'


def _enable_drag_drop(widget, on_files_dropped):
    """在 Windows 下为 widget 启用文件拖放。on_files_dropped(paths: list) 回调。"""
    if sys.platform != 'win32':
        return
    try:
        import ctypes
        from ctypes import wintypes
        shell32 = ctypes.windll.shell32
        user32 = ctypes.windll.user32
        WM_DROPFILES = 0x0233
        GWLP_WNDPROC = -4
        hwnd = widget.winfo_id()
        shell32.DragAcceptFiles(hwnd, True)
        CallWindowProcW = user32.CallWindowProcW
        CallWindowProcW.argtypes = [wintypes.LPVOID, wintypes.HWND, wintypes.UINT, wintypes.WPARAM, wintypes.LPARAM]
        CallWindowProcW.restype = wintypes.LONG_PTR
        WNDPROC = ctypes.WINFUNCTYPE(wintypes.LONG_PTR, wintypes.HWND, wintypes.UINT, wintypes.WPARAM, wintypes.LPARAM)
        old_wndproc = [None]

        def wnd_proc(hwnd, msg, wparam, lparam):
            if msg == WM_DROPFILES:
                max_len = 1024
                buf = ctypes.create_unicode_buffer(max_len)
                count = shell32.DragQueryFileW(wparam, 0xFFFFFFFF, None, 0)
                paths = []
                for i in range(count):
                    shell32.DragQueryFileW(wparam, i, buf, max_len)
                    paths.append(buf.value)
                shell32.DragFinish(wparam)
                try:
                    widget.after(0, lambda: on_files_dropped(paths))
                except Exception:
                    pass
                return 0
            if old_wndproc[0] is not None:
                return CallWindowProcW(old_wndproc[0], hwnd, msg, wparam, lparam)
            return user32.DefWindowProcW(hwnd, msg, wparam, lparam)

        new_proc = WNDPROC(wnd_proc)
        old_wndproc[0] = user32.SetWindowLongPtrW(hwnd, GWLP_WNDPROC, new_proc)
        widget._dnd_proc_ref = new_proc
    except Exception:
        pass


# 用户在小窗口输入的 DeepSeek API Key（优先于环境变量）
_stored_deepseek_api_key = None

# api.9e.lv 平台的 API Key（用于 Gemini 2.0 Flash / Gemini 3 Pro，可在 https://api.9e.lv/pricing 获取）
_stored_9e_api_key = None

# 思维链解析与流程图绘制
BOX_WIDTH = 180
BOX_HEIGHT = 52
ARROW_LEN = 28
STEP_MAX = 30
NODE_TEXT_WIDTH = 160  # 节点内文本自动换行宽度
NODE_DISPLAY_CHARS = 16  # 格子内省略显示的字符数，双击格子可看全文

# AI 外置记忆循环结构：存储路径（双层存储：调用序列 + 结点内容）
_EXTERNAL_MEMORY_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '外置记忆')
_EXTERNAL_MEMORY_FILE = os.path.join(_EXTERNAL_MEMORY_DIR, '思路存储器.json')

# 界面字体：优先使用免费可商用高级字体，按可用性回退
_UI_FONT_PREFERENCE = [
    'Source Han Sans SC', 'Source Han Sans CN', 'Noto Sans CJK SC', 'Noto Sans SC', '思源黑体',
    'Alibaba PuHuiTi', '阿里巴巴普惠体', 'HarmonyOS Sans SC',
]


def _resolve_ui_font(win):
    """解析当前可用的界面字体，返回优先列表中的第一个已安装字体。"""
    try:
        root = win.winfo_toplevel()
        families = set(f.lower() for f in tkfont.families(root))
        for name in _UI_FONT_PREFERENCE:
            if name.lower() in families:
                return name
    except Exception:
        pass
    try:
        return tkfont.nametofont('TkDefaultFont').actual()['family']
    except Exception:
        return 'sans-serif'


# ========== retrieval_label 的 label_text 存储逻辑 ==========
# 【重要】label_text 是加载外置记忆时 AI 用于检索的判据。修改存储逻辑时搜索以下函数/关键词：
# 1. 原始标签原材料：_build_retrieval_label() —— 决定从流程图哪些结点取内容，受配置 raw_parts 等控制
# 2. 格式转化：_apply_label_text_format() —— 决定最终存储格式，受配置 format_mode / ai_prompt / custom_template 控制
# 3. 写入数据库：_save_to_database() 内 raw_label、label_text 的生成与 INSERT
# 4. 配置加载/保存：_load_label_text_config()、_save_label_text_config()；默认配置 _get_default_label_text_config()
# 5. 用户可配置：主窗口「检索标签格式」按钮 -> _show_label_text_config_window()；配置文件 _LABEL_TEXT_CONFIG_FILE
# ============================================================

_LABEL_TEXT_CONFIG_FILE = os.path.join(_EXTERNAL_MEMORY_DIR, '检索标签格式配置.json')
_FANSI_PROB_CONFIG_FILE = os.path.join(_EXTERNAL_MEMORY_DIR, '反思概率配置.json')  # 引用高亮结点内容的概率 0-100

# ========== 界面风格配置（仅影响外观，不改功能） ==========
_UI_THEME_CONFIG_FILE = os.path.join(_EXTERNAL_MEMORY_DIR, '界面风格配置.json')
_current_ui_theme_id = None

_UI_THEMES = {
    'ice_cyan': {
        'name': '冰蓝晨曦', 'desc': '冰蓝·天青·浅水色，通透轻盈，高端冷调',
        'window_bg': '#e8f4f8', 'frame_bg': '#f0f9fc', 'label_bg': '#e0f2f7',
        'entry_bg': '#ffffff', 'entry_fg': '#1a3a4a', 'text_bg': '#f5fafc', 'text_fg': '#2d4a5e',
        'accent': '#0097a7', 'accent_hover': '#00acc1', 'canvas_bg': '#e3f2fd',
        'node_fill': '#b2ebf2', 'node_outline': '#00bcd4', 'node_text': '#006064',
        'node_dimmed_fill': '#e0f7fa', 'node_dimmed_outline': '#80deea', 'node_dimmed_text': '#4dd0e1',
        'placeholder': '#4dd0e1', 'font_family': 'Microsoft YaHei UI', 'font_size': 11, 'font_size_small': 9,
        'button_fg': '#ffffff', 'user_tag': '#0288d1', 'assistant_tag': '#00897b',
        'diamond_fill': '#e1f5fe', 'rounded_fill': '#e0f2f1',
    },
    'sky_teal': {
        'name': '天青琉璃', 'desc': '天空蓝·蓝绿·青瓷色，冷静克制，商务高级',
        'window_bg': '#f0f4f8', 'frame_bg': '#ffffff', 'label_bg': '#e3f2fd',
        'entry_bg': '#fafbfc', 'entry_fg': '#1e3a5f', 'text_bg': '#ffffff', 'text_fg': '#263238',
        'accent': '#00838f', 'accent_hover': '#0097a7', 'canvas_bg': '#e8eaf6',
        'node_fill': '#b3e5fc', 'node_outline': '#0288d1', 'node_text': '#01579b',
        'node_dimmed_fill': '#e1f5fe', 'node_dimmed_outline': '#90caf9', 'node_dimmed_text': '#64b5f6',
        'placeholder': '#78909c', 'font_family': 'Microsoft YaHei UI', 'font_size': 11, 'font_size_small': 9,
        'button_fg': '#ffffff', 'user_tag': '#0277bd', 'assistant_tag': '#00695c',
        'diamond_fill': '#e1f5fe', 'rounded_fill': '#e8f5e9',
    },
    'deep_ocean': {
        'name': '深海墨蓝', 'desc': '深蓝·墨青·科技感，稳重专业',
        'window_bg': '#0d1b2a', 'frame_bg': '#1b263b', 'label_bg': '#415a77',
        'entry_bg': '#2d3e50', 'entry_fg': '#e0e1dd', 'text_bg': '#1b263b', 'text_fg': '#c9d1d3',
        'accent': '#0096c7', 'accent_hover': '#00b4d8', 'canvas_bg': '#1b263b',
        'node_fill': '#2d4a6f', 'node_outline': '#48cae4', 'node_text': '#ade8f4',
        'node_dimmed_fill': '#223347', 'node_dimmed_outline': '#4a6fa5', 'node_dimmed_text': '#5c7c99',
        'placeholder': '#778da9', 'font_family': 'Microsoft YaHei UI', 'font_size': 11, 'font_size_small': 9,
        'button_fg': '#ffffff', 'user_tag': '#48cae4', 'assistant_tag': '#52b69a',
        'diamond_fill': '#2d4a6f', 'rounded_fill': '#2d4a6f',
    },
    'deep_navy': {
        'name': '深邃墨色', 'desc': '深色背景，护眼舒适，适合长时间使用',
        'window_bg': '#1a1b26', 'frame_bg': '#24283b', 'label_bg': '#1f2335',
        'entry_bg': '#414868', 'entry_fg': '#c0caf5', 'text_bg': '#1f2335', 'text_fg': '#a9b1d6',
        'accent': '#7aa2f7', 'accent_hover': '#bb9af7', 'canvas_bg': '#1f2335',
        'node_fill': '#364a7c', 'node_outline': '#7aa2f7', 'node_text': '#c0caf5',
        'node_dimmed_fill': '#2d2d44', 'node_dimmed_outline': '#565f89', 'node_dimmed_text': '#565f89',
        'placeholder': '#565f89', 'font_family': 'Microsoft YaHei UI', 'font_size': 11, 'font_size_small': 9,
        'button_fg': '#ffffff', 'user_tag': '#7aa2f7', 'assistant_tag': '#9ece6a',
        'diamond_fill': '#363b54', 'rounded_fill': '#364a7c',
    },
    'light_fresh': {
        'name': '清新浅色', 'desc': '浅色明亮，简洁清爽',
        'window_bg': '#f5f7fa', 'frame_bg': '#ffffff', 'label_bg': '#e8ecf1',
        'entry_bg': '#ffffff', 'entry_fg': '#2c3e50', 'text_bg': '#fafbfc', 'text_fg': '#2c3e50',
        'accent': '#3498db', 'accent_hover': '#2980b9', 'canvas_bg': '#f0f2f5',
        'node_fill': '#e8f4fc', 'node_outline': '#3498db', 'node_text': '#2c3e50',
        'node_dimmed_fill': '#ecf0f1', 'node_dimmed_outline': '#bdc3c7', 'node_dimmed_text': '#7f8c8d',
        'placeholder': '#95a5a6', 'font_family': 'Microsoft YaHei UI', 'font_size': 11, 'font_size_small': 9,
        'button_fg': '#ffffff', 'user_tag': '#3498db', 'assistant_tag': '#27ae60',
        'diamond_fill': '#fff5e6', 'rounded_fill': '#e0ffe0',
    },
    'warm_amber': {
        'name': '暖调舒适', 'desc': '暖色系，温和不刺眼',
        'window_bg': '#2d2a26', 'frame_bg': '#3d3832', 'label_bg': '#352f2a',
        'entry_bg': '#4a4540', 'entry_fg': '#e8dcc8', 'text_bg': '#352f2a', 'text_fg': '#d4c5a9',
        'accent': '#d4a574', 'accent_hover': '#e8c9a0', 'canvas_bg': '#352f2a',
        'node_fill': '#4a4038', 'node_outline': '#d4a574', 'node_text': '#e8dcc8',
        'node_dimmed_fill': '#3d3832', 'node_dimmed_outline': '#6b6156', 'node_dimmed_text': '#6b6156',
        'placeholder': '#6b6156', 'font_family': 'Microsoft YaHei UI', 'font_size': 11, 'font_size_small': 9,
        'button_fg': '#ffffff', 'user_tag': '#d4a574', 'assistant_tag': '#a9b665',
        'diamond_fill': '#4a4038', 'rounded_fill': '#4a4038',
    },
}


def _load_ui_theme_config():
    """加载界面风格配置，返回主题 id。"""
    try:
        if os.path.isfile(_UI_THEME_CONFIG_FILE):
            with open(_UI_THEME_CONFIG_FILE, 'r', encoding='utf-8') as f:
                cfg = json.load(f)
            tid = cfg.get('theme_id', 'deep_navy')
            if tid in _UI_THEMES:
                return tid
    except Exception:
        pass
    return 'ice_cyan'


def _save_ui_theme_config(theme_id):
    """保存界面风格配置。"""
    _ensure_memory_dir()
    try:
        with open(_UI_THEME_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump({'theme_id': theme_id}, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def _get_current_theme():
    """获取当前主题字典。"""
    global _current_ui_theme_id
    if _current_ui_theme_id is None:
        _current_ui_theme_id = _load_ui_theme_config()
    return _UI_THEMES.get(_current_ui_theme_id, _UI_THEMES['ice_cyan'])


def _get_flowchart_colors():
    """获取流程图绘制所需的颜色（由当前主题决定）。"""
    t = _get_current_theme()
    return {
        'node_fill': t.get('node_fill', '#e8f2ff'), 'node_outline': t.get('node_outline', '#2d7dff'),
        'node_text': t.get('node_text', '#1a1a1a'), 'node_dimmed_fill': t.get('node_dimmed_fill', '#d8d8d8'),
        'node_dimmed_outline': t.get('node_dimmed_outline', '#a0a0a0'),
        'node_dimmed_text': t.get('node_dimmed_text', '#707070'),
        'diamond_fill': t.get('diamond_fill', t.get('node_fill')),
        'rounded_fill': t.get('rounded_fill', t.get('node_fill')),
        'arrow': t.get('node_outline', '#2d7dff'), 'label': t.get('node_dimmed_text', '#666'),
        'loading': t.get('placeholder', '#888'),
    }


def _apply_ui_theme(win):
    """将当前主题应用到窗口及其子控件。仅外观，不改功能。"""
    global _current_ui_theme_id
    if _current_ui_theme_id is None:
        _current_ui_theme_id = _load_ui_theme_config()
    t = _get_current_theme()
    font_family = _resolve_ui_font(win)
    font_size = t.get('font_size', 11)
    font_size_small = t.get('font_size_small', 9)
    try:
        win.configure(bg=t.get('window_bg', '#1a1b26'))
    except tk.TclError:
        pass
    try:
        style = ttk.Style()
        for theme_name in ('clam', 'alt', 'default'):
            try:
                style.theme_use(theme_name)
                break
            except tk.TclError:
                pass
        style.configure('TFrame', background=t.get('frame_bg'))
        style.configure('TLabelframe', background=t.get('label_bg'))
        try:
            style.configure('TPanedwindow', background=t.get('frame_bg'))
            style.configure('TPanedwindow.Sash', width=6, background=t.get('accent', '#007acc'))
        except tk.TclError:
            pass
        style.configure('TLabelframe.Label', background=t.get('label_bg'), foreground=t.get('text_fg'), font=(font_family, font_size))
        style.configure('TButton', background=t.get('accent'), foreground=t.get('button_fg', '#ffffff'), font=(font_family, font_size), padding=(20, 10))
        style.map('TButton', background=[('active', t.get('accent_hover', t.get('accent')))])
        try:
            style.configure('TButton', relief='flat', borderwidth=1)
        except tk.TclError:
            pass
        style.configure('TLabel', background=t.get('frame_bg'), foreground=t.get('text_fg'), font=(font_family, font_size))
        style.configure('TRadiobutton', background=t.get('frame_bg'), foreground=t.get('text_fg'), font=(font_family, font_size))
        style.configure('TEntry', fieldbackground=t.get('entry_bg'), foreground=t.get('entry_fg'), insertcolor=t.get('entry_fg'), padding=4, font=(font_family, font_size))
        style.configure('TCombobox', fieldbackground=t.get('entry_bg'), foreground=t.get('entry_fg'), background=t.get('frame_bg'), font=(font_family, font_size))
        try:
            style.configure('Vertical.TScrollbar', background=t.get('accent'), troughcolor=t.get('canvas_bg', t.get('frame_bg')), arrowcolor=t.get('text_fg'), width=10)
            style.configure('Horizontal.TScrollbar', background=t.get('accent'), troughcolor=t.get('canvas_bg', t.get('frame_bg')), arrowcolor=t.get('text_fg'), width=10)
        except tk.TclError:
            pass
    except tk.TclError:
        pass

    def _recurse(w):
        try:
            if isinstance(w, tk.Canvas):
                w.configure(bg=t.get('canvas_bg'))
            elif isinstance(w, (tk.Text, scrolledtext.ScrolledText)):
                w.configure(bg=t.get('text_bg'), fg=t.get('text_fg'), insertcolor=t.get('text_fg'),
                            selectbackground=t.get('accent'), selectforeground=t.get('text_fg'),
                            font=(font_family, font_size))
                try:
                    w.tag_configure('user_tag', foreground=t.get('user_tag', '#2d7dff'))
                    w.tag_configure('assistant_tag', foreground=t.get('assistant_tag', '#0d6b0d'))
                except tk.TclError:
                    pass
            elif isinstance(w, tk.Entry):
                w.configure(bg=t.get('entry_bg'), fg=t.get('entry_fg'), insertbackground=t.get('entry_fg'),
                            font=(font_family, font_size))
        except tk.TclError:
            pass
        for c in w.winfo_children():
            _recurse(c)
    try:
        _recurse(win)
    except tk.TclError:
        pass


def _show_ui_theme_selector(parent, on_theme_changed=None):
    """打开界面风格选择窗口。on_theme_changed: 可选，切换后回调（用于重绘流程图等）。"""
    top = tk.Toplevel(parent)
    top.title('界面风格')
    top.geometry('480x420')
    top.transient(parent)
    t = _get_current_theme()
    try:
        top.configure(bg=t.get('window_bg'))
    except tk.TclError:
        pass
    f = ttk.Frame(top, padding=16)
    f.pack(fill=tk.BOTH, expand=True)
    ttk.Label(f, text='选择界面风格（保存后立即生效）').pack(anchor=tk.W)
    var = tk.StringVar(value=_current_ui_theme_id or _load_ui_theme_config())
    for tid, theme in _UI_THEMES.items():
        rb = ttk.Radiobutton(f, text='%s — %s' % (theme['name'], theme['desc']), variable=var, value=tid)
        rb.pack(anchor=tk.W, pady=4)
    def on_ok():
        global _current_ui_theme_id
        tid = var.get()
        if tid not in _UI_THEMES:
            return
        _current_ui_theme_id = tid
        _save_ui_theme_config(tid)
        _apply_ui_theme(parent)
        if on_theme_changed:
            on_theme_changed()
        try:
            top.destroy()
        except tk.TclError:
            pass
        messagebox.showinfo('界面风格', '已保存，新风格已应用。', parent=parent)
    ttk.Button(f, text='保存并应用', command=on_ok, width=14).pack(anchor=tk.W, pady=(12, 0))
    _apply_ui_theme(top)


# 内置 SQLite 数据库（零配置，首次使用时自动创建，与外置记忆同目录）
_SQLITE_DB_PATH = os.path.join(_EXTERNAL_MEMORY_DIR, 'ai_memory.db')

# MySQL 配置（可选，用于需要共享数据库的场景；若使用需配合 Wampserver 等，并手动建库建表）
_DB_CONFIG = {
    'host': '127.0.0.1',
    'port': 3306,
    'user': 'root',
    'password': '',
    'database': 'ai_memory',
    'charset': 'utf8mb4',
}


def _get_sqlite_conn():
    """获取 SQLite 连接，不存在时自动创建数据库和表。"""
    _ensure_memory_dir()
    conn = sqlite3.connect(_SQLITE_DB_PATH)
    conn.row_factory = lambda c, r: dict(zip([col[0] for col in c.description], r))
    cur = conn.cursor()
    cur.executescript('''
        CREATE TABLE IF NOT EXISTS flowchart_content (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            content TEXT,
            node_type VARCHAR(50)
        );
        CREATE TABLE IF NOT EXISTS flowchart_session (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            mode VARCHAR(50),
            model_name VARCHAR(100),
            summary VARCHAR(500),
            node_sequence TEXT
        );
        CREATE TABLE IF NOT EXISTS retrieval_label (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id INTEGER,
            label_text TEXT
        );
    ''')
    conn.commit()
    return conn


def _get_default_label_text_config():
    """label_text 存储格式的默认配置。"""
    return {
        'raw_parts': 'after_first_and_before_last',  # after_first_and_before_last | all | first_only | last_only
        'after_first_limit': 300,
        'before_last_limit': 300,
        'separator': ' | ',
        'format_mode': 'ai',  # ai | raw | custom
        'ai_prompt': '请根据以下思维链内容，用一句话概括为「开头内容、最终结果、最终目的」的格式。\n要求：输出仅此一句话，不要其他解释。格式示例：开头讨论X问题，最终得出Y结论，目的是Z。\n\n思维链内容：\n{raw_label}',
        'custom_template': '{raw_label}',
        'output_max_len': 500,
        'retrieval_timeout_seconds': 30,  # AI 检索超时(秒)，超时后回退到字符串检索
    }


def _load_label_text_config():
    """加载 label_text 配置，失败则返回默认配置。"""
    try:
        if os.path.isfile(_LABEL_TEXT_CONFIG_FILE):
            with open(_LABEL_TEXT_CONFIG_FILE, 'r', encoding='utf-8') as f:
                cfg = json.load(f)
            default = _get_default_label_text_config()
            for k in default:
                if k not in cfg:
                    cfg[k] = default[k]
            return cfg
    except Exception:
        pass
    return _get_default_label_text_config()


def _save_label_text_config(cfg):
    """保存 label_text 配置到文件。"""
    _ensure_memory_dir()
    try:
        with open(_LABEL_TEXT_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def _ensure_memory_dir():
    """确保外置记忆目录存在。"""
    if not os.path.exists(_EXTERNAL_MEMORY_DIR):
        try:
            os.makedirs(_EXTERNAL_MEMORY_DIR, exist_ok=True)
        except OSError:
            pass


def _load_fansi_prob():
    """加载反思时引用高亮结点内容的概率（0-100）。默认 50。"""
    try:
        if os.path.exists(_FANSI_PROB_CONFIG_FILE):
            with open(_FANSI_PROB_CONFIG_FILE, 'r', encoding='utf-8') as f:
                d = json.load(f)
                p = int(d.get('prob', 50))
                return max(0, min(100, p))
    except Exception:
        pass
    return 50


def _save_fansi_prob(prob):
    """保存反思概率配置。"""
    _ensure_memory_dir()
    try:
        with open(_FANSI_PROB_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump({'prob': max(0, min(100, int(prob)))}, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def _save_external_memory(flow_steps, flow_spec, messages):
    """
    AI 外置记忆循环结构：将迭代过程信息有机存储于外置存储器。
    双层存储：调用序列（思路骨架）+ 结点内容（思路结点存储器）。
    数据格式便于后续模型微调或训练。
    """
    _ensure_memory_dir()
    data = {
        'version': 1,
        'updated_at': datetime.now().isoformat(),
        'call_sequence': [],
        'nodes': [],
        'edges': [],
        'flow_steps': list(flow_steps) if flow_steps else [],
        'messages': list(messages) if messages else [],
    }
    if flow_spec and flow_spec.get('nodes'):
        nodes = flow_spec['nodes']
        data['call_sequence'] = [n.get('id', i + 1) for i, n in enumerate(nodes)]
        data['nodes'] = [
            {'id': n.get('id', i + 1), 'type': n.get('type', 'rect'), 'text': str(n.get('text', '')), 'version': 'v1'}
            for i, n in enumerate(nodes)
        ]
        data['edges'] = list(flow_spec.get('edges', []))
    else:
        data['call_sequence'] = list(range(1, len(flow_steps) + 1))
        data['nodes'] = [
            {'id': i + 1, 'type': 'rect', 'text': str(s), 'version': 'v1'}
            for i, s in enumerate(flow_steps or [])
        ]
    try:
        with io.open(_EXTERNAL_MEMORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except (IOError, OSError, TypeError):
        return False


def _load_external_memory():
    """从外置存储器加载思路数据。成功返回 (flow_steps, flow_spec, messages)，失败返回 None。"""
    if not os.path.isfile(_EXTERNAL_MEMORY_FILE):
        return None
    try:
        with io.open(_EXTERNAL_MEMORY_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except (IOError, OSError, ValueError, TypeError):
        return None
    nodes = data.get('nodes', [])
    edges = data.get('edges', [])
    flow_steps = data.get('flow_steps', [])
    messages = data.get('messages', [])
    if not nodes and not flow_steps:
        return None
    flow_spec = None
    if nodes:
        flow_spec = {'nodes': nodes, 'edges': edges}
        if not flow_steps:
            flow_steps = [str(n.get('text', '')) for n in nodes]
    return flow_steps, flow_spec, messages


def _refine_last_node_by_interaction(flow_steps, flow_spec, num_bright, interaction_content):
    """
    单思路结点动态更新：通过交互反馈，将 AI 的回复补充到当前（最后高亮）思路结点。
    「AI 能够通过交互反馈时事更新当前回答思路」。
    当事件结束后，思路栈内信息由外置记忆保留。
    """
    if not interaction_content or not str(interaction_content).strip():
        return False
    content = str(interaction_content).strip()[:2000]
    if not content:
        return False
    updated = False
    if flow_spec and flow_spec.get('nodes') and num_bright > 0:
        idx = min(num_bright, len(flow_spec['nodes'])) - 1
        if idx >= 0:
            node = flow_spec['nodes'][idx]
            old = str(node.get('text', '')).strip()
            node['text'] = old + '\n\n[交互更新] ' + content
            node.pop('db_content_id', None)  # 内容已变更，保存时需重新 get_or_insert
            updated = True
    if flow_steps and num_bright > 0:
        idx = min(num_bright, len(flow_steps)) - 1
        if idx >= 0:
            flow_steps[idx] = str(flow_steps[idx]).strip() + '\n\n[交互更新] ' + content
            updated = True
    return updated


def _is_flowchart_fully_bright(flow_steps, flow_spec, num_bright):
    """判断思维链流程图是否全亮（高亮节点数 >= 总节点数）。"""
    total = len(flow_spec['nodes']) if (flow_spec and flow_spec.get('nodes')) else len(flow_steps or [])
    return total > 0 and num_bright >= total


def _build_retrieval_label(nodes_data):
    """
    生成检索标签原材料，根据配置决定从流程图哪些结点取内容。
    nodes_data: [(content, node_type, db_content_id?), ...]，使用 content 即 d[0]。
    配置项 raw_parts 可选：after_first_and_before_last | all | first_only | last_only
    """
    cfg = _load_label_text_config()
    if not nodes_data:
        return ''
    contents = [str(d[0] or '') for d in nodes_data]
    parts = []
    raw_parts = cfg.get('raw_parts', 'after_first_and_before_last')
    after_limit = int(cfg.get('after_first_limit', 300))
    before_limit = int(cfg.get('before_last_limit', 300))
    sep = cfg.get('separator', ' | ')
    if len(contents) == 1:
        return contents[0][:500]
    if raw_parts == 'first_only':
        parts = [contents[0][:after_limit]]
    elif raw_parts == 'last_only':
        parts = [contents[-1][:before_limit]]
    elif raw_parts == 'all':
        parts = [''.join(contents)[:after_limit + before_limit]]
    else:
        after_first = ''.join(contents[1:])[:after_limit]
        before_last = ''.join(contents[:-1])[:before_limit]
        parts = [after_first, before_last] if (after_first or before_last) else [contents[0][:300]]
    return (sep.join(p for p in parts if p.strip())).strip() or contents[0][:300]


def _apply_label_text_format(raw_label, mode, ollama_model):
    """
    根据配置将原始标签 raw_label 转为最终存储的 label_text。
    format_mode: ai=调用模型格式化 | raw=直接使用 | custom=使用模板替换 {raw_label}
    """
    cfg = _load_label_text_config()
    max_len = int(cfg.get('output_max_len', 500))
    if not raw_label or not str(raw_label).strip():
        return (raw_label or '')[:max_len]
    raw_label = str(raw_label).strip()
    fmt_mode = cfg.get('format_mode', 'ai')
    if fmt_mode == 'raw':
        return raw_label[:max_len]
    if fmt_mode == 'custom':
        tpl = cfg.get('custom_template', '{raw_label}')
        return tpl.replace('{raw_label}', raw_label)[:max_len]
    prompt = (cfg.get('ai_prompt', '') or '思维链内容：\n{raw_label}').replace('{raw_label}', raw_label[:1500])
    msgs = [{'role': 'user', 'content': prompt}]
    try:
        if mode == 'ollama' and ollama_model:
            body, _ = call_ollama_api(msgs, ollama_model, use_think=False)
        else:
            body, _ = _call_cloud_api(msgs, mode)
        formatted = (body or '').strip()
        return formatted[:max_len] if formatted else raw_label[:max_len]
    except Exception:
        return raw_label[:max_len]


def _show_label_text_config_window(parent):
    """打开检索标签格式设置窗口，用于配置 label_text 的存储格式与参考结点位置。"""
    top = tk.Toplevel(parent)
    top.title('检索标签格式设置')
    top.geometry('560x480')
    top.transient(parent)
    cfg = _load_label_text_config()

    f = ttk.Frame(top, padding=12)
    f.pack(fill=tk.BOTH, expand=True)

    raw_parts_options = (
        'after_first_and_before_last - 开头结点后+结束结点前（默认）',
        'all - 全部结点内容',
        'first_only - 仅第一个结点',
        'last_only - 仅最后一个结点',
    )
    rp_map = {'after_first_and_before_last': raw_parts_options[0], 'all': raw_parts_options[1],
              'first_only': raw_parts_options[2], 'last_only': raw_parts_options[3]}
    ttk.Label(f, text='参考流程图结点位置（raw_parts）：').pack(anchor=tk.W)
    raw_parts_var = tk.StringVar(value=rp_map.get(cfg.get('raw_parts'), raw_parts_options[0]))
    ttk.Combobox(f, textvariable=raw_parts_var, width=45, state='readonly', values=raw_parts_options).pack(fill=tk.X, pady=(2, 8))

    lim_row = ttk.Frame(f)
    lim_row.pack(fill=tk.X, pady=(0, 8))
    ttk.Label(lim_row, text='开头后取字数：').pack(side=tk.LEFT, padx=(0, 4))
    after_limit_var = tk.StringVar(value=str(cfg.get('after_first_limit', 300)))
    ttk.Entry(lim_row, textvariable=after_limit_var, width=8).pack(side=tk.LEFT, padx=(0, 16))
    ttk.Label(lim_row, text='结束前取字数：').pack(side=tk.LEFT, padx=(0, 4))
    before_limit_var = tk.StringVar(value=str(cfg.get('before_last_limit', 300)))
    ttk.Entry(lim_row, textvariable=before_limit_var, width=8).pack(side=tk.LEFT)

    fmt_options = ('ai - AI 格式化', 'raw - 直接使用原始拼接', 'custom - 自定义模板')
    fmt_map = {'ai': fmt_options[0], 'raw': fmt_options[1], 'custom': fmt_options[2]}
    ttk.Label(f, text='格式模式（format_mode）：').pack(anchor=tk.W, pady=(8, 0))
    format_var = tk.StringVar(value=fmt_map.get(cfg.get('format_mode'), fmt_options[0]))
    ttk.Combobox(f, textvariable=format_var, width=30, state='readonly', values=fmt_options).pack(fill=tk.X, pady=(2, 4))

    ttk.Label(f, text='AI 格式化提示词（format_mode=ai 时使用，{raw_label} 为占位符）：').pack(anchor=tk.W, pady=(8, 0))
    ai_prompt_text = scrolledtext.ScrolledText(f, height=5, width=60, wrap=tk.WORD)
    ai_prompt_text.pack(fill=tk.X, pady=(2, 4))
    ai_prompt_text.insert(tk.END, cfg.get('ai_prompt', ''))

    ttk.Label(f, text='自定义模板（format_mode=custom 时使用，{raw_label} 为占位符）：').pack(anchor=tk.W, pady=(8, 0))
    custom_var = tk.StringVar(value=cfg.get('custom_template', '{raw_label}'))
    ttk.Entry(f, textvariable=custom_var, width=60).pack(fill=tk.X, pady=(2, 4))

    ttk.Label(f, text='输出最大长度（output_max_len）：').pack(anchor=tk.W, pady=(4, 0))
    max_len_var = tk.StringVar(value=str(cfg.get('output_max_len', 500)))
    ttk.Entry(f, textvariable=max_len_var, width=10).pack(anchor=tk.W, pady=(2, 4))

    ttk.Label(f, text='AI 检索超时(秒)（retrieval_timeout_seconds）：超时后回退到字符串检索').pack(anchor=tk.W, pady=(8, 0))
    timeout_var = tk.StringVar(value=str(cfg.get('retrieval_timeout_seconds', 30)))
    ttk.Entry(f, textvariable=timeout_var, width=10).pack(anchor=tk.W, pady=(2, 4))

    def raw_parts_value():
        v = raw_parts_var.get()
        if 'after_first' in v:
            return 'after_first_and_before_last'
        if 'all' in v:
            return 'all'
        if 'first_only' in v:
            return 'first_only'
        if 'last_only' in v:
            return 'last_only'
        return 'after_first_and_before_last'

    def fmt_value():
        v = format_var.get()
        if 'raw' in v:
            return 'raw'
        if 'custom' in v:
            return 'custom'
        return 'ai'

    def on_save():
        try:
            new_cfg = dict(cfg)
            new_cfg['raw_parts'] = raw_parts_value()
            new_cfg['after_first_limit'] = int(after_limit_var.get() or 300)
            new_cfg['before_last_limit'] = int(before_limit_var.get() or 300)
            new_cfg['format_mode'] = fmt_value()
            new_cfg['ai_prompt'] = ai_prompt_text.get(1.0, tk.END).strip() or _get_default_label_text_config()['ai_prompt']
            new_cfg['custom_template'] = custom_var.get().strip() or '{raw_label}'
            new_cfg['output_max_len'] = int(max_len_var.get() or 500)
            new_cfg['retrieval_timeout_seconds'] = max(5, int(timeout_var.get() or 30))
            if _save_label_text_config(new_cfg):
                messagebox.showinfo('保存', '检索标签格式配置已保存。', parent=top)
                top.destroy()
            else:
                messagebox.showerror('保存', '保存失败，请检查目录权限。', parent=top)
        except ValueError as e:
            messagebox.showerror('格式错误', '请检查数字字段是否正确：%s' % e, parent=top)

    ttk.Button(f, text='保存', command=on_save).pack(anchor=tk.W, pady=(12, 0))


def _get_or_insert_content_id(cur, conn, content, node_type):
    """获取或插入 flowchart_content，返回 id。旧结点不重复存储。支持 SQLite / MySQL。"""
    content = str(content or '')
    node_type = str(node_type or 'rect')[:50]
    cur.execute(
        'SELECT id FROM flowchart_content WHERE content = ? AND node_type = ? LIMIT 1',
        (content, node_type)
    )
    row = cur.fetchone()
    if row:
        return row['id'] if isinstance(row, dict) else row[0]
    cur.execute(
        'INSERT INTO flowchart_content (content, node_type) VALUES (?, ?)',
        (content, node_type)
    )
    return cur.lastrowid


def _save_to_database(flow_steps, flow_spec, mode, model_name, summary='', ollama_model=None):
    """
    将全亮流程图保存到内置 SQLite 数据库。label_text 先经 AI 按「开头内容、最终结果、最终目的」格式生成后再存储。
    结点去重：从数据库加载的结点保留 db_content_id，直接复用，减少 SELECT；新结点才 get_or_insert。
    """
    try:
        conn = _get_sqlite_conn()
        cur = conn.cursor()
        nodes_data = []
        if flow_spec and flow_spec.get('nodes'):
            for n in flow_spec['nodes']:
                nodes_data.append((
                    str(n.get('text', '')),
                    n.get('type', 'rect') or 'rect',
                    n.get('db_content_id')
                ))
        elif flow_steps:
            for s in flow_steps:
                nodes_data.append((str(s), 'rect', None))
        if not nodes_data:
            conn.close()
            return False
        node_ids = []
        for content, ntype, db_content_id in nodes_data:
            if db_content_id is not None:
                node_ids.append(int(db_content_id))
            else:
                nid = _get_or_insert_content_id(cur, conn, content, ntype)
                node_ids.append(nid)
        node_sequence = json.dumps(node_ids)
        cur.execute(
            'INSERT INTO flowchart_session (mode, model_name, summary, node_sequence) VALUES (?, ?, ?, ?)',
            (mode or 'unknown', model_name or '', summary[:500] if summary else '', node_sequence)
        )
        session_id = cur.lastrowid
        raw_label = _build_retrieval_label(nodes_data)
        label_text = _apply_label_text_format(
            raw_label, mode,
            ollama_model if mode == 'ollama' else None
        )
        cur.execute(
            'INSERT INTO retrieval_label (session_id, label_text) VALUES (?, ?)',
            (session_id, label_text or raw_label[:500])
        )
        conn.commit()
        conn.close()
        return True
    except Exception:
        try:
            conn.rollback()
            conn.close()
        except Exception:
            pass
        return False


def _string_search_retrieval_label(rows, query):
    """
    字符串检索：在 label_text 中查找包含查询内容的记录，返回 session_id。
    优先完全包含，否则按子串匹配长度选取最佳。
    """
    if not query or not rows:
        return 0
    q = query.strip()
    if not q:
        return 0
    q_lower = q.lower()
    best_sid = 0
    best_len = 0
    for r in rows:
        label = (r.get('label_text') or '').strip().lower()
        if not label:
            continue
        if q_lower in label:
            return r['session_id']
        for ln in range(min(len(q), len(label)), 0, -1):
            for i in range(len(q) - ln + 1):
                sub = q_lower[i:i + ln]
                if sub in label and ln > best_len:
                    best_len = ln
                    best_sid = r['session_id']
                    break
    return best_sid


def _load_from_database_by_query(parent, user_query, mode, ollama_model, on_loaded):
    """
    根据用户输入的检索描述，用 AI 匹配 retrieval_label，找到对应 flowchart_session，
    再按 node_sequence 从 flowchart_content 取结点，加载外置记忆。
    若 AI 检索超时，则回退到字符串检索；仍无匹配则提示。
    使用内置 SQLite，零配置。
    on_loaded: 成功时回调 (flow_steps, flow_spec)，可异步。
    """
    if not user_query or not str(user_query).strip():
        parent.after(0, lambda: messagebox.showinfo('加载', '请输入检索描述', parent=parent))
        return
    query = str(user_query).strip()

    def run():
        try:
            conn = _get_sqlite_conn()
            cur = conn.cursor()
            cur.execute('SELECT id, session_id, label_text FROM retrieval_label ORDER BY id ASC')
            rows = cur.fetchall()
            conn.close()
        except Exception as e:
            parent.after(0, lambda: messagebox.showerror('加载', '无法读取数据库：%s' % e, parent=parent))
            return
        if not rows:
            parent.after(0, lambda: messagebox.showinfo('加载', '数据库中没有可检索的外置记忆', parent=parent))
            return
        labels_text = '\n'.join([
            'id=%d, session_id=%d: %s' % (r['id'], r['session_id'], (r.get('label_text') or '').strip())
            for r in rows
        ])
        prompt = '''你负责匹配用户要加载的思维链流程图。以下为 retrieval_label 表中**全部**检索标签，每个 id 对应一条 label_text，session_id 对应要加载的流程图。

你必须完整阅读所有 id 对应的 label_text 内容后，再做出选择。

全部检索标签：
%s

用户输入的检索描述：%s

请从上述全部选项中选出与用户描述最匹配的一个 session_id。只回复数字（session_id），不要任何其他文字。若无匹配可回复 0。''' % (labels_text, query[:500])
        msgs = [{'role': 'user', 'content': prompt}]
        session_id = 0
        timeout_sec = max(5, int(_load_label_text_config().get('retrieval_timeout_seconds', 30)))

        def do_ai_match():
            if mode == 'ollama' and ollama_model:
                body, _ = call_ollama_api(msgs, ollama_model, use_think=False)
            else:
                body, _ = _call_cloud_api(msgs, mode)
            body = (body or '').strip()
            for w in re.findall(r'\d+', body):
                sid = int(w)
                if sid in (r['session_id'] for r in rows):
                    return sid
            return 0

        try:
            with ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(do_ai_match)
                session_id = fut.result(timeout=timeout_sec)
        except FuturesTimeoutError:
            session_id = _string_search_retrieval_label(rows, query)
        except Exception:
            session_id = _string_search_retrieval_label(rows, query)
        if session_id == 0:
            parent.after(0, lambda: messagebox.showinfo('加载', '未找到匹配的流程图', parent=parent))
            return
        try:
            conn = _get_sqlite_conn()
            cur = conn.cursor()
            cur.execute('SELECT node_sequence FROM flowchart_session WHERE id = ?', (session_id,))
            row = cur.fetchone()
            conn.close()
            if not row:
                parent.after(0, lambda: messagebox.showinfo('加载', '会话不存在', parent=parent))
                return
            node_seq = row['node_sequence'] if isinstance(row, dict) else row[0]
            try:
                node_ids = json.loads(node_seq)
            except (TypeError, ValueError):
                parent.after(0, lambda: messagebox.showerror('加载', '数据格式错误', parent=parent))
                return
            conn = _get_sqlite_conn()
            cur = conn.cursor()
            flow_steps = []
            flow_spec_nodes = []
            for i, nid in enumerate(node_ids):
                cur.execute('SELECT content, node_type FROM flowchart_content WHERE id = ?', (nid,))
                nr = cur.fetchone()
                if nr:
                    content = nr['content'] if isinstance(nr, dict) else nr[0]
                    ntype = (nr['node_type'] if isinstance(nr, dict) else nr[1]) or 'rect'
                    flow_steps.append(content)
                    flow_spec_nodes.append({'id': i + 1, 'type': ntype, 'text': content, 'db_content_id': nid})
            conn.close()
            flow_spec = {'nodes': flow_spec_nodes, 'edges': [{'from': i, 'to': i + 1, 'label': ''} for i in range(1, len(flow_spec_nodes))]}
        except Exception as e:
            parent.after(0, lambda: messagebox.showerror('加载', '无法加载：%s' % e, parent=parent))
            return
        parent.after(0, lambda: on_loaded(flow_steps, flow_spec))
    threading.Thread(target=run, daemon=True).start()


def fangfa1(reasoning_content):
    """
    从 DeepSeek 返回的 reasoning_content 中解析出思维链步骤列表，
    用于后续绘制横向流程图。
    若 reasoning_content 为空或 None，返回空列表。
    """
    if not reasoning_content or not str(reasoning_content).strip():
        return []
    text = str(reasoning_content).strip()
    steps = []
    # 先按双换行分段落
    parts = re.split(r'\n\s*\n', text)
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # 再按单换行分，避免一个段落过长
        lines = p.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 按句号、问号、感叹号分句（保留短句合并）
            for sent in re.split(r'(?<=[。.!?])\s+', line):
                sent = sent.strip()
                if sent:
                    steps.append(sent)
        if len(steps) >= STEP_MAX:
            break
    # 若没有按句号分出多步，则把较长段落切成多步（按长度）
    if len(steps) <= 1 and parts:
        for p in parts:
            p = p.strip()
            if not p:
                continue
            chunk_size = 80
            for i in range(0, len(p), chunk_size):
                steps.append(p[i:i + chunk_size])
                if len(steps) >= STEP_MAX:
                    break
            if len(steps) >= STEP_MAX:
                break
    return steps[:STEP_MAX]


def _abbrev_text(text, max_chars=NODE_DISPLAY_CHARS):
    """省略显示：超过 max_chars 时截断并加 …"""
    if not text:
        return ''
    s = str(text).strip().replace('\n', ' ')
    return (s[:max_chars] + '…') if len(s) > max_chars else s


def _wrap_text(text, width_chars):
    """将长文本按字符数折行。"""
    if not text:
        return ''
    lines, cur, cur_len = [], [], 0
    for c in text:
        cur.append(c)
        cur_len += 1
        if c in '。！？\n' or cur_len >= width_chars:
            lines.append(''.join(cur))
            cur, cur_len = [], 0
    if cur:
        lines.append(''.join(cur))
    return '\n'.join(lines)


def _extract_flowchart_json_from_content(content):
    """
    从模型回复中提取流程图 JSON。约定格式：---思维链流程图JSON---\\n{...}\\n---END---
    成功返回 dict，否则返回 None。支持嵌套 JSON 的括号匹配。
    """
    if not content or not isinstance(content, str):
        return None
    start_marker, end_marker = '---思维链流程图JSON---', '---END---'
    if start_marker not in content or end_marker not in content:
        return None
    start_idx = content.find(start_marker) + len(start_marker)
    end_idx = content.find(end_marker)
    block = content[start_idx:end_idx].strip()
    i = block.find('{')
    if i < 0:
        return None
    depth, j = 0, i
    for j, c in enumerate(block[i:], i):
        if c == '{':
            depth += 1
        elif c == '}':
            depth -= 1
            if depth == 0:
                try:
                    parsed = json.loads(block[i:j + 1])
                    if parsed.get('nodes') and parsed.get('edges') is not None:
                        return parsed
                except json.JSONDecodeError:
                    pass
                return None
    return None


def _content_without_flowchart_block(content):
    """移除回复中的流程图 JSON 块，返回纯文本用于显示。"""
    if not content:
        return content
    return re.sub(r'\s*---思维链流程图JSON---[\s\S]*?---END---\s*', '', str(content)).strip()


def _filter_redundant_nodes(new_nodes, bright_texts):
    """过滤与已有高亮节点内容冗余的新节点（子串重复或完全相同）。"""
    if not bright_texts:
        return new_nodes
    bright_list = [str(t).strip().lower() for t in bright_texts if t and str(t).strip()]
    filtered = []
    for n in new_nodes:
        t = str(n.get('text', '')).strip()
        if not t:
            continue
        t_lower = t.lower()
        is_redundant = any(b in t_lower or t_lower in b for b in bright_list)
        if not is_redundant:
            filtered.append(n)
    return filtered if filtered else new_nodes


def _generate_flowchart_spec(reasoning_content, mode, ollama_model=None):
    """
    调用模型将思维链内容转化为专业流程图规范。
    返回 {"nodes": [{"id":1,"type":"rect","text":"..."}], "edges": [{"from":1,"to":2,"label":""}]}
    type: rect(矩形流程), diamond(菱形判断), rounded(圆角开始结束)
    """
    prompt = f'''你是指南式流程图专家。将下面的思维链内容转化为标准流程图结构的JSON。
输出必须是纯JSON，不要任何其他文字、markdown或说明。
格式：{{"nodes":[{{"id":1,"type":"rect","text":"步骤内容"}}],"edges":[{{"from":1,"to":2,"label":"是"}}]}}
type 说明：rect=矩形(普通流程步骤)，diamond=菱形(判断/分支)，rounded=圆角矩形(开始或结束)。
判断节点(diamond)可有多个出边，用 label 区分如"是"/"否"。
节点 id 从 1 开始递增。edges 的 from/to 为节点 id。
思维链内容：
{reasoning_content[:6000]}'''
    msgs = [{'role': 'user', 'content': prompt}]
    try:
        if mode == 'ollama' and ollama_model:
            body, _ = call_ollama_api(msgs, ollama_model, use_think=False)
        else:
            body, _ = _call_cloud_api(msgs, mode)
        j = json.loads(re.search(r'\{[\s\S]*\}', body or '{}').group(0))
        if j.get('nodes') and j.get('edges') is not None:
            return j
    except Exception:
        pass
    return None


def _interactive_zoom(canvas, flow_zoom_ref, flow_steps, flow_spec, num_bright, event, delta_override=None):
    """交互模式窗口的 flowchart 缩放。"""
    delta = delta_override if delta_override is not None else (event.delta if hasattr(event, 'delta') else 0)
    if delta == 0:
        return
    flow_zoom_ref[0] *= (1.1 if delta > 0 else 0.9)
    flow_zoom_ref[0] = max(0.5, min(2.0, flow_zoom_ref[0]))
    _draw_flowchart(canvas, flow_steps, flow_spec, flow_zoom_ref[0], num_bright)


def _clear_placeholder_interactive(entry):
    """交互窗口输入框占位符清除。"""
    if entry.get() == '在此输入问题，按 Enter 发送':
        entry.delete(0, tk.END)


def _show_api_key_dialog(parent):
    """双击 DeepSeek API（云端）时弹出的小窗口，用于输入 API Key。"""
    global _stored_deepseek_api_key
    top = tk.Toplevel(parent)
    top.title('设置 DeepSeek API Key')
    top.geometry('420x180')
    top.minsize(380, 140)
    top.transient(parent)
    top.grab_set()
    f = ttk.Frame(top, padding=12)
    f.pack(fill=tk.BOTH, expand=True)
    ttk.Label(f, text='请输入 DeepSeek API Key（可在 platform.deepseek.com 获取）：').pack(anchor=tk.W)
    key_var = tk.StringVar(value=_stored_deepseek_api_key or '')
    entry = ttk.Entry(f, textvariable=key_var, width=52, show='')
    entry.pack(fill=tk.X, pady=(6, 12))
    entry.focus_set()

    def on_confirm():
        global _stored_deepseek_api_key
        key = key_var.get().strip()
        if not key:
            messagebox.showwarning('提示', '请输入 API Key', parent=top)
            return
        if key in ('你的API_Key', '你的API Key', 'your_api_key'):
            messagebox.showwarning('提示', '请不要使用占位符，请输入真实的 API Key', parent=top)
            return
        try:
            key.encode('ascii')
        except UnicodeEncodeError:
            messagebox.showwarning('提示', 'API Key 只能包含英文和数字', parent=top)
            return
        _stored_deepseek_api_key = key
        try:
            top.destroy()
        except tk.TclError:
            pass
        messagebox.showinfo('设置成功', 'API Key 已保存，可以调用 DeepSeek 了。', parent=parent)

    def on_cancel():
        try:
            top.destroy()
        except tk.TclError:
            pass

    btn_frame = ttk.Frame(f)
    btn_frame.pack(fill=tk.X)
    ttk.Button(btn_frame, text='确认', command=on_confirm, width=10).pack(side=tk.LEFT, padx=(0, 8))
    ttk.Button(btn_frame, text='取消', command=on_cancel, width=10).pack(side=tk.LEFT)
    top.bind('<Return>', lambda e: on_confirm())
    top.bind('<Escape>', lambda e: on_cancel())
    _apply_ui_theme(top)


def _on_liangshao(flow_canvas, flow_steps, flow_spec, flow_zoom, num_bright_ref, redraw_flow, update_continue_btn_state):
    """亮少：当前高亮结点中的最后一个不再高亮。"""
    if num_bright_ref[0] <= 0:
        return
    num_bright_ref[0] -= 1
    redraw_flow()
    update_continue_btn_state()


def _on_caiqu(flow_canvas, flow_steps, flow_spec, flow_zoom, num_bright_ref, parent, redraw_flow, update_continue_btn_state):
    """裁去：弹出窗口让用户选择从第几个高亮结点裁去（该结点及之后全部去掉）。"""
    nb = num_bright_ref[0]
    if nb <= 0:
        messagebox.showinfo('裁去', '当前无高亮结点，无法裁去。', parent=parent)
        return
    total = len(flow_spec['nodes']) if (flow_spec and flow_spec.get('nodes')) else len(flow_steps or [])
    if total <= 0:
        return
    # 高亮结点序号为 1..nb
    top = tk.Toplevel(parent)
    top.title('裁去 - 选择起始结点')
    top.geometry('380x200')
    top.transient(parent)
    f = ttk.Frame(top, padding=12)
    f.pack(fill=tk.BOTH, expand=True)
    ttk.Label(f, text='选择从第几个高亮结点裁去（该结点及之后全部删除）：').pack(anchor=tk.W)
    options = [str(i) for i in range(1, nb + 1)]
    var = tk.StringVar(value=options[0])
    cb = ttk.Combobox(f, textvariable=var, values=options, state='readonly', width=8)
    cb.pack(anchor=tk.W, pady=8)
    cb.current(0)
    hint = '高亮结点序号：' + '、'.join(str(i) for i in options)
    ttk.Label(f, text=hint).pack(anchor=tk.W)
    result = [None]

    def on_confirm():
        try:
            sel = int(var.get())
            if 1 <= sel <= nb:
                result[0] = sel
                top.destroy()
        except (ValueError, tk.TclError):
            pass

    def on_cancel():
        try:
            top.destroy()
        except tk.TclError:
            pass

    btn_f = ttk.Frame(f)
    btn_f.pack(fill=tk.X, pady=(12, 0))
    ttk.Button(btn_f, text='确定', command=on_confirm).pack(side=tk.LEFT, padx=(0, 8))
    ttk.Button(btn_f, text='取消', command=on_cancel).pack(side=tk.LEFT)
    top.bind('<Return>', lambda e: on_confirm())
    top.bind('<Escape>', lambda e: on_cancel())
    _apply_ui_theme(top)
    top.grab_set()
    parent.wait_window(top)

    sel = result[0]
    if sel is None or sel < 1:
        return
    cut_from_idx = sel - 1  # 0-based，保留 0..cut_from_idx-1，删除 cut_from_idx 及之后
    new_len = cut_from_idx
    if new_len <= 0:
        flow_steps.clear()
        if flow_spec and flow_spec.get('nodes'):
            flow_spec['nodes'].clear()
            flow_spec['edges'] = flow_spec.get('edges', [])
            flow_spec['edges'].clear()
    else:
        flow_steps[:] = flow_steps[:new_len]
        if flow_spec and flow_spec.get('nodes'):
            flow_spec['nodes'][:] = flow_spec['nodes'][:new_len]
            keep_ids = {n.get('id', i + 1) for i, n in enumerate(flow_spec['nodes'])}
            flow_spec['edges'] = [e for e in flow_spec.get('edges', []) if e.get('from') in keep_ids and e.get('to') in keep_ids]
    num_bright_ref[0] = min(nb, new_len)
    redraw_flow()
    update_continue_btn_state()


def _on_fansi(flow_steps, flow_spec, num_bright_ref, chat_text, flow_canvas, flow_zoom, parent, app_self, status_var, send_btn, append_chat, update_continue_btn_state):
    """反思：根据当前会话内容，去掉未高亮结点，重新生成最后一个高亮结点之后的新思维链。"""
    nb = num_bright_ref[0]
    total = len(flow_spec['nodes']) if (flow_spec and flow_spec.get('nodes')) else len(flow_steps or [])
    if total <= 0 or nb <= 0:
        messagebox.showinfo('反思', '流程图为空或无高亮结点。', parent=parent)
        return
    session_content = chat_text.get(1.0, tk.END).strip()
    if not session_content:
        messagebox.showinfo('反思', '会话窗口为空。', parent=parent)
        return
    bright_text = _get_bright_node_texts(flow_steps, flow_spec, nb)

    # 弹出设置窗口：概率可调，保存后执行
    top = tk.Toplevel(parent)
    top.title('反思 - 概率设置')
    top.geometry('360x180')
    top.transient(parent)
    f = ttk.Frame(top, padding=12)
    f.pack(fill=tk.BOTH, expand=True)
    ttk.Label(f, text='引用高亮结点内容的概率（0–100%）：').pack(anchor=tk.W)
    ttk.Label(f, text='满足该概率时，API 续写会引用高亮结点；否则仅用会话内容。').pack(anchor=tk.W)
    prob_var = tk.StringVar(value=str(_load_fansi_prob()))
    entry = ttk.Entry(f, textvariable=prob_var, width=8)
    entry.pack(anchor=tk.W, pady=8)
    ttk.Label(f, text='0–100，执行时按此概率决定是否引用高亮结点内容。').pack(anchor=tk.W)
    do_run = [False]
    prob_saved = [_load_fansi_prob()]

    def on_save():
        try:
            p = int(prob_var.get().strip())
            p = max(0, min(100, p))
            prob_var.set(str(p))
            if _save_fansi_prob(p):
                prob_saved[0] = p
                messagebox.showinfo('反思', '概率已保存。', parent=top)
        except (ValueError, tk.TclError):
            messagebox.showwarning('反思', '请输入 0–100 的整数。', parent=top)

    def on_execute():
        on_save()  # 执行前先保存当前值
        do_run[0] = True
        try:
            top.destroy()
        except tk.TclError:
            pass

    def on_cancel():
        try:
            top.destroy()
        except tk.TclError:
            pass

    btn_f = ttk.Frame(f)
    btn_f.pack(fill=tk.X, pady=(12, 0))
    ttk.Button(btn_f, text='保存', command=on_save).pack(side=tk.LEFT, padx=(0, 8))
    ttk.Button(btn_f, text='执行反思', command=on_execute).pack(side=tk.LEFT, padx=(0, 8))
    ttk.Button(btn_f, text='取消', command=on_cancel).pack(side=tk.LEFT)
    top.bind('<Return>', lambda e: on_execute())
    top.bind('<Escape>', lambda e: on_cancel())
    _apply_ui_theme(top)
    top.grab_set()
    parent.wait_window(top)

    if not do_run[0]:
        return

    use_bright = random.random() < (prob_saved[0] / 100.0)
    if not bright_text.strip() and use_bright:
        messagebox.showinfo('反思', '高亮结点内容为空。', parent=parent)
        return
    # 先裁去未高亮结点
    flow_steps[:] = flow_steps[:nb]
    if flow_spec and flow_spec.get('nodes') and nb < len(flow_spec['nodes']):
        flow_spec['nodes'][:] = flow_spec['nodes'][:nb]
        keep_ids = {n.get('id', i + 1) for i, n in enumerate(flow_spec['nodes'])}
        flow_spec['edges'] = [e for e in flow_spec.get('edges', []) if e.get('from') in keep_ids and e.get('to') in keep_ids]
    status_var.set('正在反思并续写思维链…')
    send_btn.configure(state=tk.DISABLED)

    def run():
        try:
            mode = app_self.mode_var.get()
            ollama_model = app_self.ollama_model_var.get().strip() if mode == 'ollama' else None
            if use_bright and bright_text.strip():
                prompt_body = f'''当前会话内容：
{session_content[:3000]}

当前思维链（高亮部分）：
{bright_text}'''
            else:
                prompt_body = f'''当前会话内容：
{session_content[:3000]}

（本次不引用高亮结点内容，仅基于会话内容续写）'''
            fc_prompt = f'''基于以下内容，请续写思维链（在已有步骤之后添加新结点）。输出必须是纯流程图JSON，不要其他文字。

【重要】续写内容必须与已有思维链不重复。专业流程图规则：rect=流程步骤(长方形)，diamond=判断分支(菱形)，rounded=开始/结束(圆角)。

在回答末尾，请严格按以下格式输出续写的专业流程图JSON：
---思维链流程图JSON---
{{"nodes":[{{"id":1,"type":"rect","text":"新步骤"}}],"edges":[{{"from":1,"to":2,"label":""}}]}}
---END---

{prompt_body}'''
            fc_msgs = [{'role': 'user', 'content': fc_prompt}]
            if mode == 'ollama':
                fc_content, fc_reasoning = call_ollama_api(fc_msgs, ollama_model, use_think=False)
            else:
                fc_content, fc_reasoning = _call_cloud_api(fc_msgs, mode)
            new_spec = _extract_flowchart_json_from_content(fc_content) or _extract_flowchart_json_from_content(fc_reasoning)
            if not new_spec and fc_reasoning and len(fc_reasoning.strip()) > 50:
                new_spec = _generate_flowchart_spec(fc_reasoning, mode, ollama_model)
            if new_spec and new_spec.get('nodes'):
                bright_texts = [str(n.get('text', '')) for n in flow_spec['nodes'][:nb]] if (flow_spec and flow_spec.get('nodes')) else list(flow_steps[:nb])
                new_spec['nodes'] = _filter_redundant_nodes(new_spec['nodes'], bright_texts)
            if new_spec and new_spec.get('nodes'):
                if not flow_spec.get('nodes') and flow_steps:
                    flow_spec['nodes'] = [{'id': i + 1, 'type': 'rect', 'text': s} for i, s in enumerate(flow_steps)]
                    flow_spec['edges'] = [{'from': i, 'to': i + 1, 'label': ''} for i in range(1, len(flow_steps))]
                nodes = flow_spec['nodes']
                last_id = max(n.get('id', i + 1) for i, n in enumerate(nodes))
                base_id = last_id + 1
                new_nodes = new_spec['nodes']
                id_map = {n.get('id', i + 1): base_id + i for i, n in enumerate(new_nodes)}
                flow_spec.setdefault('edges', []).append({'from': last_id, 'to': base_id, 'label': ''})
                for i, n in enumerate(new_nodes):
                    flow_spec['nodes'].append({'id': base_id + i, 'type': n.get('type', 'rect'), 'text': n.get('text', '')})
                for e in new_spec.get('edges', []):
                    f, t = e.get('from'), e.get('to')
                    if f in id_map and t in id_map:
                        flow_spec['edges'].append({'from': id_map[f], 'to': id_map[t], 'label': e.get('label', '')})
                for n in new_nodes:
                    flow_steps.append(str(n.get('text', '')))
                parent.after(0, lambda: _draw_flowchart(flow_canvas, flow_steps, flow_spec, flow_zoom[0], nb))
            else:
                parent.after(0, lambda: messagebox.showinfo('反思', '未能解析到新的思维链结点。', parent=parent))
        except Exception as e:
            parent.after(0, lambda: messagebox.showerror('反思失败', str(e), parent=parent))
        finally:
            parent.after(0, lambda: (send_btn.configure(state=tk.NORMAL), status_var.set('就绪'), update_continue_btn_state()))
            parent.after(0, lambda: _save_external_memory(flow_steps, flow_spec, app_self.messages))

    threading.Thread(target=run, daemon=True).start()


def _interactive_node_click(event, canvas, flow_steps, flow_spec, parent):
    """交互模式窗口的双击节点弹窗。"""
    ids = canvas.find_closest(event.x, event.y)
    if not ids:
        return
    tags = canvas.gettags(ids[0])
    node_tag = next((t for t in tags if t.startswith('node_')), None)
    if not node_tag:
        return
    try:
        idx = int(node_tag.split('_')[1])
    except (ValueError, IndexError):
        return
    full_text = ''
    if flow_spec and flow_spec.get('nodes'):
        nodes = flow_spec['nodes']
        if 0 <= idx < len(nodes):
            full_text = str(nodes[idx].get('text', ''))
    elif flow_steps and 0 <= idx < len(flow_steps):
        full_text = flow_steps[idx]
    if not full_text:
        return
    top = tk.Toplevel(parent)
    top.title('思维链节点 - 完整内容')
    top.geometry('500x360')
    top.minsize(400, 200)
    btn_frame = ttk.Frame(top)
    btn_frame.pack(fill=tk.X, padx=8, pady=4)
    ttk.Button(btn_frame, text='关闭', command=top.destroy, width=28).pack(side=tk.RIGHT)
    txt = tk.Text(top, wrap=tk.WORD, font=('Microsoft YaHei UI', 10), padx=8, pady=8)
    txt.pack(fill=tk.BOTH, expand=True)
    txt.insert(tk.END, full_text)
    txt.configure(state=tk.DISABLED)


def _get_bright_node_texts(steps, flow_spec, num_bright):
    """获取前 num_bright 个节点的文本内容，拼接成字符串。"""
    if num_bright <= 0:
        return ''
    texts = []
    if flow_spec and flow_spec.get('nodes'):
        nodes = flow_spec['nodes']
        for i in range(min(num_bright, len(nodes))):
            texts.append(str(nodes[i].get('text', '')))
    elif steps:
        for i in range(min(num_bright, len(steps))):
            texts.append(str(steps[i]))
    return '\n\n'.join(t for t in texts if t.strip())


def _extract_flowchart_as_text(flow_steps, flow_spec):
    """
    将思维链流程图内容提取为可选择复制的字符串，保留格式并标注节点形状。
    若有内容返回字符串，否则返回 None。
    """
    TYPE_LABELS = {
        'rect': '矩形(流程步骤)',
        'diamond': '菱形(判断)',
        'rounded': '圆角(开始/结束)',
    }
    lines = []
    if flow_spec and flow_spec.get('nodes'):
        nodes = flow_spec['nodes']
        edges = flow_spec.get('edges', [])
        for i, n in enumerate(nodes):
            nid = n.get('id', i + 1)
            ntype = (n.get('type') or 'rect').lower()
            label = TYPE_LABELS.get(ntype, ntype)
            text = str(n.get('text', '')).strip()
            if text:
                lines.append('[节点%d - %s]\n%s' % (nid, label, _wrap_text(text, NODE_TEXT_WIDTH)))
        if edges:
            lines.append('')
            lines.append('--- 连接关系 ---')
            for e in edges:
                fid, tid = e.get('from'), e.get('to')
                lbl = e.get('label', '')
                lbl_str = (' [%s]' % lbl) if lbl else ''
                lines.append('节点%d → 节点%d%s' % (fid, tid, lbl_str))
    elif flow_steps:
        for i, step in enumerate(flow_steps):
            text = str(step).strip()
            if text:
                lines.append('[节点%d - 矩形(流程步骤)]\n%s' % (i + 1, _wrap_text(text, NODE_TEXT_WIDTH)))
    if not lines:
        return None
    return '\n\n'.join(lines)


def _get_flowchart_nodes_ordered(flow_steps, flow_spec):
    """
    按顺序返回 [(节点文本, 节点功能标签), ...]。
    用于循环自思考时逐节点执行。
    """
    TYPE_LABELS = {
        'rect': '矩形(流程步骤)',
        'diamond': '菱形(判断)',
        'rounded': '圆角(开始/结束)',
    }
    result = []
    if flow_spec and flow_spec.get('nodes'):
        for i, n in enumerate(flow_spec['nodes']):
            ntype = (n.get('type') or 'rect').lower()
            label = TYPE_LABELS.get(ntype, ntype)
            text = str(n.get('text', '')).strip()
            result.append((text, label))
    elif flow_steps:
        for i, step in enumerate(flow_steps):
            text = str(step).strip()
            result.append((text, '矩形(流程步骤)'))
    return result


def _show_zisikao_mode_dialog(parent, app_self):
    """点击自思考后：选择直接自思考或循环自思考。"""
    has_flow = len(app_self.flow_steps) > 0 or (app_self.flow_spec and app_self.flow_spec.get('nodes'))
    if not has_flow:
        messagebox.showinfo('自思考', '思维链流程图为空，请先进行对话获取思维链后再使用。', parent=parent)
        return
    top = tk.Toplevel(parent)
    top.title('自思考 - 选择模式')
    top.geometry('380x180')
    top.resizable(False, False)
    ttk.Label(top, text='请选择自思考模式：', font=(_resolve_ui_font(top), 10)).pack(pady=(16, 12))
    f = ttk.Frame(top)
    f.pack(pady=8, padx=16, fill=tk.X)
    ttk.Button(f, text='直接自思考', width=14,
               command=lambda: (top.destroy(), app_self._do_zisikao_direct())).pack(side=tk.LEFT, padx=(0, 12), ipady=6)
    ttk.Button(f, text='循环自思考', width=14,
               command=lambda: (top.destroy(), app_self._do_zisikao_loop())).pack(side=tk.LEFT, ipady=6)
    ttk.Label(top, text='直接：流程图+会话→模型→一次返回\n循环：按节点顺序逐步执行，每步输出作为下一步输入',
              font=(_resolve_ui_font(top), 9), justify=tk.LEFT).pack(pady=(8, 0), padx=16, anchor=tk.W)
    _apply_ui_theme(top)


def _show_extract_window(parent, flow_steps, flow_spec):
    """点击提取按钮时：若有内容则在新窗口显示可复制的文本，否则提示无法提取。"""
    extracted = _extract_flowchart_as_text(flow_steps, flow_spec)
    if not extracted:
        messagebox.showinfo('提取', '思维链流程图为空，无法提取。')
        return
    top = tk.Toplevel(parent)
    top.title('思维链流程图 - 提取内容')
    top.geometry('600x500')
    top.minsize(450, 300)
    btn_frame = ttk.Frame(top)
    btn_frame.pack(fill=tk.X, padx=8, pady=4)
    ttk.Button(btn_frame, text='关闭', command=top.destroy, width=12).pack(side=tk.RIGHT)
    txt = tk.Text(top, wrap=tk.WORD, font=('Microsoft YaHei UI', 10), padx=8, pady=8, state=tk.NORMAL)
    txt.pack(fill=tk.BOTH, expand=True)
    txt.insert(tk.END, extracted)


def _draw_flowchart(canvas, steps, flow_spec=None, zoom=1.0, num_bright=None, loading=False):
    """
    绘制思维链流程图。flow_spec 为模型生成的专业规范；否则用 steps 简单线性绘制。
    num_bright: 若提供，前 num_bright 个节点保持原亮度，其余变暗（交互模式用）。
    loading: 若为 True 且无内容，显示「正在生成专业流程图…」。
    """
    canvas.delete('all')
    font_family = _resolve_ui_font(canvas)
    scale = max(0.5, min(2.0, zoom))
    cw = max(canvas.winfo_reqwidth(), 400)
    ch = max(canvas.winfo_reqheight(), 80)
    if loading and not steps and not (flow_spec and flow_spec.get('nodes')):
        fc = _get_flowchart_colors()
        canvas.create_text(cw // 2, ch // 2, text='正在生成专业流程图…',
                          fill=fc.get('loading', '#888'), font=(font_family, 10))
        canvas.configure(scrollregion=(0, 0, cw, ch))
        return

    if flow_spec and flow_spec.get('nodes'):
        nodes = flow_spec['nodes']
        edges = flow_spec.get('edges', [])
        padding = int(24 * scale)
        bw = int(BOX_WIDTH * scale)
        bh = int(BOX_HEIGHT * scale)
        font_sz = max(8, int(10 * scale))
        cols = min(6, max(1, int((cw - 2 * padding) / (bw + 20))))
        positions = {}
        for i, n in enumerate(nodes):
            r, c = i // cols, i % cols
            x = padding + c * (bw + 28)
            y = padding + r * (bh + 36)
            positions[n.get('id', i + 1)] = (x + bw // 2, y + bh // 2, bw, bh)

        for i, n in enumerate(nodes):
            nid = n.get('id') or (i + 1)
            ntype = (n.get('type') or 'rect').lower()
            full_text = str(n.get('text', ''))
            display_text = _abbrev_text(full_text)
            if nid not in positions:
                continue
            cx, cy, bw, bh = positions[nid]
            x1, y1 = cx - bw // 2, cy - bh // 2
            x2, y2 = cx + bw // 2, cy + bh // 2
            fc = _get_flowchart_colors()
            dimmed = num_bright is not None and i >= num_bright
            if dimmed:
                fill, outline, tfill = fc['node_dimmed_fill'], fc['node_dimmed_outline'], fc['node_dimmed_text']
            else:
                fill, outline, tfill = fc['node_fill'], fc['node_outline'], fc['node_text']
            tag = 'node_%d' % i
            if ntype == 'diamond' and not dimmed:
                canvas.create_polygon([cx, y1, x2, cy, cx, y2, x1, cy], outline=outline, fill=fc.get('diamond_fill', fill), width=2, tags=(tag,))
            elif ntype == 'rounded' and not dimmed:
                canvas.create_rectangle(x1, y1, x2, y2, outline=outline, fill=fc.get('rounded_fill', fill), width=2, tags=(tag,))
            else:
                canvas.create_rectangle(x1, y1, x2, y2, outline=outline, fill=fill, width=2, tags=(tag,))
            canvas.create_text(cx, cy, text=display_text, fill=tfill, font=(font_family, font_sz),
                              width=int(NODE_TEXT_WIDTH * scale), tags=(tag,))

        fc = _get_flowchart_colors()
        for e in edges:
            fid, tid = e.get('from'), e.get('to')
            if fid in positions and tid in positions:
                x1, y1 = positions[fid][0], positions[fid][1]
                x2, y2 = positions[tid][0], positions[tid][1]
                canvas.create_line(x1, y1, x2, y2, fill=fc.get('arrow', '#2d7dff'), width=2, arrow=tk.LAST)
                if e.get('label'):
                    mx, my = (x1 + x2) / 2, (y1 + y2) / 2
                    canvas.create_text(mx, my - 8, text=str(e['label']), fill=fc.get('label', '#666'), font=(font_family, 8))

        max_x = max(p[0] + p[2] // 2 for p in positions.values()) if positions else cw
        max_y = max(p[1] + p[3] // 2 for p in positions.values()) if positions else ch
        canvas.configure(scrollregion=(0, 0, max_x + padding, max_y + padding))
        return

    if not steps:
        fc = _get_flowchart_colors()
        canvas.create_text(cw // 2, ch // 2, text='暂无思维链（或未使用 reasoner 模型）',
                          fill=fc.get('loading', '#888'), font=(font_family, 10))
        canvas.configure(scrollregion=(0, 0, cw, ch))
        return

    fc = _get_flowchart_colors()
    padding = int(20 * scale)
    bw = int(BOX_WIDTH * scale)
    bh = int(BOX_HEIGHT * scale)
    font_sz = max(8, int(10 * scale))
    x = padding
    for i, step in enumerate(steps):
        display_text = _abbrev_text(step)
        tag = 'node_%d' % i
        dimmed = num_bright is not None and i >= num_bright
        if dimmed:
            outline, fill, tfill = fc['node_dimmed_outline'], fc['node_dimmed_fill'], fc['node_dimmed_text']
        else:
            outline, fill, tfill = fc['node_outline'], fc['node_fill'], fc['node_text']
        canvas.create_rectangle(x, padding, x + bw, padding + bh,
                                outline=outline, fill=fill, width=2, tags=(tag,))
        canvas.create_text(x + bw // 2, padding + bh // 2, text=display_text, fill=tfill,
                          font=(font_family, font_sz), width=int(NODE_TEXT_WIDTH * scale), tags=(tag,))
        x += bw + int(ARROW_LEN * scale)
        if i < len(steps) - 1:
            canvas.create_line(x - int(ARROW_LEN * scale), padding + bh // 2,
                              x, padding + bh // 2, fill=fc.get('arrow', '#2d7dff'), width=2, arrow=tk.LAST)
    canvas.configure(scrollregion=(0, 0, x + padding, padding + bh + 24))


def call_deepseek_api(messages):
    """
    直接使用 requests 调用 DeepSeek API，避免 openai/httpx 在 Windows 下的 ascii 编码问题。
    返回 (content, reasoning_content) 或抛出异常。
    优先使用窗口输入的 API Key，其次使用环境变量。
    """
    global _stored_deepseek_api_key
    if requests is None:
        raise RuntimeError('请先安装: pip install requests')
    api_key = (_stored_deepseek_api_key or os.environ.get('DEEPSEEK_API_KEY') or os.environ.get('OPENAI_API_KEY') or '').strip()
    if not api_key:
        raise ValueError(
            '请先设置 API Key。\n'
            '方式一：双击窗口左上角「DeepSeek API（云端）」可弹出输入窗口。\n'
            '方式二：设置环境变量 DEEPSEEK_API_KEY（在 platform.deepseek.com 获取）'
        )
    if api_key in ('你的API_Key', '你的API Key', 'your_api_key'):
        raise ValueError('请将「你的API_Key」替换成你在 platform.deepseek.com 获取的真实 API Key，不要使用占位符')
    try:
        api_key.encode('ascii')
    except UnicodeEncodeError:
        raise ValueError('API Key 只能包含英文和数字，不能包含中文。请设置正确的 DEEPSEEK_API_KEY（格式通常为 sk- 开头）')
    url = 'https://api.deepseek.com/v1/chat/completions'
    headers = {
        'Content-Type': 'application/json; charset=utf-8',
        'Authorization': 'Bearer ' + api_key,
    }
    payload = {
        'model': 'deepseek-reasoner',
        'messages': messages,
        'max_tokens': 8192,
    }
    # 显式使用 UTF-8 编码，避免 ascii 编码错误
    body = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    try:
        resp = requests.post(url, headers=headers, data=body, timeout=120)
        resp.raise_for_status()
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 401:
            raise ValueError(
                '401 认证失败：API Key 无效或错误。\n'
                '请到 https://platform.deepseek.com/api_keys 检查并重新获取有效 Key，\n'
                '确保 CMD 中已正确执行：set DEEPSEEK_API_KEY=你的真实Key'
            )
        if e.response.status_code == 402:
            raise ValueError(
                '402 余额不足：请到 https://platform.deepseek.com/top_up 为账户充值。'
            )
        raise
    data = resp.json()
    choice = data.get('choices', [{}])[0]
    msg = choice.get('message', {})
    content = (msg.get('content') or '').strip()
    reasoning = msg.get('reasoning_content') or ''
    return content, reasoning


# api.9e.lv 平台（Gemini 2.0 Flash / Gemini 3 Pro，OpenAI 兼容接口）
API_9E_BASE = 'https://api.9e.lv'


def call_9e_api(messages, model):
    """
    调用 api.9e.lv 的 OpenAI 兼容接口。
    model: 'gemini-2.0-flash' 或 'gemini-3-pro-preview'
    返回 (content, reasoning_content)。
    """
    global _stored_9e_api_key
    if requests is None:
        raise RuntimeError('请先安装: pip install requests')
    api_key = (_stored_9e_api_key or os.environ.get('API_9E_KEY') or os.environ.get('OPENAI_API_KEY') or '').strip()
    if not api_key:
        raise ValueError(
            '请先设置 api.9e.lv 的 API Key。\n'
            '方式一：双击「Gemini 2.0 Flash」或「Gemini 3 Pro」可弹出输入窗口。\n'
            '方式二：设置环境变量 API_9E_KEY\n'
            'Key 可在 https://api.9e.lv/pricing 获取'
        )
    if api_key in ('你的API_Key', '你的API Key', 'your_api_key'):
        raise ValueError('请输入真实的 API Key，不要使用占位符')
    try:
        api_key.encode('ascii')
    except UnicodeEncodeError:
        raise ValueError('API Key 只能包含英文和数字')
    url = API_9E_BASE + '/v1/chat/completions'
    headers = {
        'Content-Type': 'application/json; charset=utf-8',
        'Authorization': 'Bearer ' + api_key,
    }
    payload = {
        'model': model,
        'messages': messages,
        'max_tokens': 8192,
    }
    body = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    try:
        resp = requests.post(url, headers=headers, data=body, timeout=120)
        resp.raise_for_status()
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 401:
            raise ValueError('401 认证失败：API Key 无效。请到 https://api.9e.lv/pricing 检查 Key。')
        if e.response.status_code == 402:
            raise ValueError('402 余额不足：请为 api.9e.lv 账户充值。')
        raise
    data = resp.json()
    choice = data.get('choices', [{}])[0]
    msg = choice.get('message', {})
    content = (msg.get('content') or '').strip()
    reasoning = msg.get('reasoning_content') or ''
    return content, reasoning


def _show_9e_api_key_dialog(parent):
    """双击 Gemini（9e.lv）时弹出的小窗口，用于输入 API Key。"""
    global _stored_9e_api_key
    top = tk.Toplevel(parent)
    top.title('设置 api.9e.lv API Key')
    top.geometry('460x180')
    top.minsize(400, 140)
    top.transient(parent)
    top.grab_set()
    f = ttk.Frame(top, padding=12)
    f.pack(fill=tk.BOTH, expand=True)
    ttk.Label(f, text='请输入 api.9e.lv 的 API Key（可在 https://api.9e.lv/pricing 获取）：').pack(anchor=tk.W)
    key_var = tk.StringVar(value=_stored_9e_api_key or '')
    entry = ttk.Entry(f, textvariable=key_var, width=52, show='')
    entry.pack(fill=tk.X, pady=(6, 12))
    entry.focus_set()

    def on_confirm():
        global _stored_9e_api_key
        key = key_var.get().strip()
        if not key:
            messagebox.showwarning('提示', '请输入 API Key', parent=top)
            return
        if key in ('你的API_Key', '你的API Key', 'your_api_key'):
            messagebox.showwarning('提示', '请不要使用占位符，请输入真实的 API Key', parent=top)
            return
        try:
            key.encode('ascii')
        except UnicodeEncodeError:
            messagebox.showwarning('提示', 'API Key 只能包含英文和数字', parent=top)
            return
        _stored_9e_api_key = key
        try:
            top.destroy()
        except tk.TclError:
            pass
        messagebox.showinfo('设置成功', 'API Key 已保存，可以调用 Gemini 模型了。', parent=parent)

    def on_cancel():
        try:
            top.destroy()
        except tk.TclError:
            pass

    btn_frame = ttk.Frame(f)
    btn_frame.pack(fill=tk.X)
    ttk.Button(btn_frame, text='确认', command=on_confirm, width=10).pack(side=tk.LEFT, padx=(0, 8))
    ttk.Button(btn_frame, text='取消', command=on_cancel, width=10).pack(side=tk.LEFT)
    top.bind('<Return>', lambda e: on_confirm())
    top.bind('<Escape>', lambda e: on_cancel())
    _apply_ui_theme(top)


def _call_cloud_api(messages, mode):
    """根据 mode 调用对应的云端 API，返回 (content, reasoning)。"""
    if mode == 'deepseek':
        return call_deepseek_api(messages)
    if mode == 'gemini_flash':
        return call_9e_api(messages, 'gemini-2.0-flash')
    if mode == 'gemini_pro':
        return call_9e_api(messages, 'gemini-3-pro-preview')
    return None, None


def _get_cloud_assistant_name(mode):
    """获取云端模式下的助手显示名。"""
    if mode == 'deepseek':
        return 'DeepSeek'
    if mode == 'gemini_flash':
        return 'Gemini 2.0 Flash'
    if mode == 'gemini_pro':
        return 'Gemini 3 Pro'
    return 'DeepSeek'


# Ollama 本地模型
OLLAMA_BASE = 'http://localhost:11434'


def fetch_ollama_models():
    """获取 Ollama 已部署的模型列表，返回 [模型名, ...]。"""
    if requests is None:
        raise RuntimeError('请先安装: pip install requests')
    try:
        resp = requests.get(f'{OLLAMA_BASE}/api/tags', timeout=5)
        resp.raise_for_status()
        data = resp.json()
        models = data.get('models', [])
        return [m.get('name', '') for m in models if m.get('name')]
    except requests.exceptions.ConnectionError:
        raise ValueError('无法连接 Ollama。请确保 Ollama 已启动（运行 ollama serve 或启动 Ollama 应用）。')
    except Exception as e:
        raise ValueError(f'Ollama 模型列表获取失败：{e}')


OLLAMA_MAX_CONTEXT = 12  # 多轮对话时最多保留最近消息数，避免请求过长导致 400


def _extract_think_from_content(content):
    """从 content 中提取 <think>...</think>` 或 <thinking>...</thinking> 块内容，用于思维链流程图。"""
    if not content:
        return ''
    parts = []
    for pattern in (r'<think>(.*?)</think>', r'<thinking>(.*?)</thinking>'):
        parts.extend(re.findall(pattern, content, re.DOTALL | re.IGNORECASE))
    return '\n\n'.join(p.strip() for p in parts if p.strip())


def _sanitize_for_ollama(s):
    """移除可能引起 Ollama 400 的控制字符。"""
    if not s:
        return ' '
    s = str(s).strip() or ' '
    return ''.join(c for c in s if c == '\n' or c == '\t' or c == '\r' or ord(c) >= 32)


def _ollama_messages_to_api(messages):
    """将消息列表转为 Ollama 所需的格式。支持多模态：含 images 的消息会保留 images 数组。"""
    msgs = []
    for m in messages:
        role = m.get('role', 'user')
        content = m.get('content', '')
        if isinstance(content, list):
            # 云端多模态格式：提取文本部分
            text_parts = [p.get('text', '') for p in content if isinstance(p, dict) and p.get('type') == 'text']
            content = _sanitize_for_ollama(' '.join(text_parts))
        else:
            content = _sanitize_for_ollama(str(content))
        item = {'role': role, 'content': content}
        if 'images' in m and m['images']:
            item['images'] = m['images']
        msgs.append(item)
    if len(msgs) > OLLAMA_MAX_CONTEXT:
        msgs = msgs[-OLLAMA_MAX_CONTEXT:]
    return msgs


def call_ollama_api(messages, model, use_think=False):
    """
    调用本地 Ollama API。支持思维链的模型（如 deepseek-r1）可传 use_think=True。
    返回 (content, reasoning_content)。
    """
    if requests is None:
        raise RuntimeError('请先安装: pip install requests')
    if not model:
        raise ValueError('请先在 Ollama 模式下选择模型')
    url = f'{OLLAMA_BASE}/api/chat'
    api_messages = _ollama_messages_to_api(messages)

    def _do_request(include_think):
        p = {'model': model, 'messages': api_messages, 'stream': False}
        if include_think:
            p['think'] = True
        body = json.dumps(p, ensure_ascii=False).encode('utf-8')
        return requests.post(
            url,
            headers={'Content-Type': 'application/json'},
            data=body,
            timeout=300,
        )

    try:
        resp = _do_request(include_think=use_think)
        if resp.status_code == 400 and use_think:
            resp = _do_request(include_think=False)
        resp.raise_for_status()
    except requests.exceptions.ConnectionError:
        raise ValueError('无法连接 Ollama。请确保 Ollama 已启动（运行 ollama serve）。')
    except requests.exceptions.HTTPError as e:
        err_detail = ''
        try:
            j = e.response.json()
            err_detail = j.get('error', '')
        except Exception:
            pass
        if not err_detail and e.response.text:
            err_detail = e.response.text[:200]
        msg = f'{e.response.status_code} {e.response.reason}'
        if err_detail:
            msg += f'\n详情：{err_detail}'
        raise ValueError(f'Ollama 请求失败：{msg}')
    data = resp.json()
    msg = data.get('message', {})
    content = (msg.get('content') or '').strip()
    reasoning = msg.get('thinking') or msg.get('reasoning_content') or ''
    if not reasoning:
        reasoning = _extract_think_from_content(content)
    return content, reasoning


class DeepSeekChatApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title('DeepSeek 思维链对话')
        self.root.geometry('1400x900')
        self.root.minsize(1000, 700)

        self.messages = []
        self.thinking_steps = []
        self.shen = 1  # 交互模式窗口计数，关闭窗口时恢复为1
        self._flowchart_loading = False  # 专业流程图生成中，不显示简单流程图
        self._flowchart_refresh_paused = False  # 暂停刷新：禁止自动更新流程图
        self.mode_var = tk.StringVar(value='deepseek')  # 'deepseek' | 'ollama'
        self.ollama_model_var = tk.StringVar(value='')
        self._building_ui()
        self._client = None
        if os.path.isfile(_EXTERNAL_MEMORY_FILE):
            self.status_var.set('检测到外置记忆，可点击「加载外置记忆」恢复')

    def _building_ui(self):
        # 模式选择与模型选择
        ctrl_frame = ttk.Frame(self.root, padding=(12, 12, 12, 0))
        ctrl_frame.pack(fill=tk.X)
        self.deepseek_rb = ttk.Radiobutton(ctrl_frame, text='DeepSeek API（云端）', variable=self.mode_var, value='deepseek',
                        command=self._on_mode_change)
        self.deepseek_rb.pack(side=tk.LEFT, padx=(0, 16))
        self.deepseek_rb.bind('<Double-Button-1>', lambda e: _show_api_key_dialog(self.root))
        gemini_flash_rb = ttk.Radiobutton(ctrl_frame, text='Gemini 2.0 Flash', variable=self.mode_var, value='gemini_flash',
                        command=self._on_mode_change)
        gemini_flash_rb.pack(side=tk.LEFT, padx=(0, 12))
        gemini_flash_rb.bind('<Double-Button-1>', lambda e: _show_9e_api_key_dialog(self.root))
        gemini_pro_rb = ttk.Radiobutton(ctrl_frame, text='Gemini 3 Pro', variable=self.mode_var, value='gemini_pro',
                        command=self._on_mode_change)
        gemini_pro_rb.pack(side=tk.LEFT, padx=(0, 12))
        gemini_pro_rb.bind('<Double-Button-1>', lambda e: _show_9e_api_key_dialog(self.root))
        ttk.Radiobutton(ctrl_frame, text='Ollama 本地', variable=self.mode_var, value='ollama',
                        command=self._on_mode_change).pack(side=tk.LEFT, padx=(0, 8))

        self.ollama_frame = ttk.Frame(ctrl_frame)
        self.ollama_frame.pack(side=tk.LEFT, padx=(8, 0))
        ttk.Label(self.ollama_frame, text='模型：').pack(side=tk.LEFT, padx=(0, 4))
        self.model_combo = ttk.Combobox(
            self.ollama_frame, textvariable=self.ollama_model_var,
            width=28
        )
        self.model_combo.pack(side=tk.LEFT, padx=(0, 4))
        ttk.Button(self.ollama_frame, text='刷新模型', command=self._refresh_ollama_models).pack(side=tk.LEFT)
        self._on_mode_change()

        # 使用 PanedWindow 支持手动拖动调整思维链/对话区域高度
        main_paned = ttk.PanedWindow(self.root, orient=tk.VERTICAL)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=8, pady=4)

        flowchart_frame = ttk.LabelFrame(main_paned, text=' 思维链流程图（点击后滚动鼠标缩放） ', padding=8)
        main_paned.add(flowchart_frame, weight=0)

        flow_btn_row = ttk.Frame(flowchart_frame)
        flow_btn_row.pack(fill=tk.X, pady=(0, 8))
        ttk.Button(flow_btn_row, text='保存外置记忆', command=self._on_save_memory).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(flow_btn_row, text='加载外置记忆', command=self._on_load_memory).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(flow_btn_row, text='检索标签格式', command=lambda: _show_label_text_config_window(self.root)).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(flow_btn_row, text='界面风格', command=lambda: _show_ui_theme_selector(self.root, on_theme_changed=self._redraw_flowchart)).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(flow_btn_row, text='清空', command=self._on_clear_all).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(flow_btn_row, text='提取', command=lambda: _show_extract_window(self.root, self.flow_steps, self.flow_spec)).pack(side=tk.RIGHT, padx=(0, 10))
        ttk.Button(flow_btn_row, text='自思考', command=lambda: _show_zisikao_mode_dialog(self.root, self)).pack(side=tk.RIGHT, padx=(0, 10))
        self.pause_refresh_btn = ttk.Button(flow_btn_row, text='暂停刷新', command=self._on_pause_refresh)
        self.pause_refresh_btn.pack(side=tk.RIGHT, padx=(0, 10))
        self.resume_refresh_btn = ttk.Button(flow_btn_row, text='继续刷新', command=self._on_resume_refresh, state=tk.DISABLED)
        self.resume_refresh_btn.pack(side=tk.RIGHT)

        canvas_container = ttk.Frame(flowchart_frame)
        canvas_container.pack(fill=tk.BOTH, expand=True)

        self.flow_canvas = tk.Canvas(
            canvas_container,
            bg='#f5f5f5',
            highlightthickness=1,
            highlightbackground='#ccc'
        )
        h_scroll = ttk.Scrollbar(canvas_container, orient=tk.HORIZONTAL, command=self.flow_canvas.xview)
        v_scroll = ttk.Scrollbar(canvas_container, orient=tk.VERTICAL, command=self.flow_canvas.yview)

        self.flow_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.flow_canvas.configure(xscrollcommand=h_scroll.set, yscrollcommand=v_scroll.set)
        self.flow_zoom = 1.0
        self.flow_steps = []
        self.flow_spec = None
        self.flow_canvas.bind('<MouseWheel>', self._on_flowchart_zoom)
        self.flow_canvas.bind('<Button-4>', lambda e: self._on_flowchart_zoom(e, 120))
        self.flow_canvas.bind('<Button-5>', lambda e: self._on_flowchart_zoom(e, -120))
        self.flow_canvas.bind('<Enter>', lambda e: self.flow_canvas.focus_set())
        flowchart_frame.bind('<Configure>', lambda e: self._redraw_flowchart())
        self.flow_canvas.bind('<Double-Button-1>', self._on_flowchart_node_double_click)

        _draw_flowchart(self.flow_canvas, [], zoom=self.flow_zoom)

        chat_frame = ttk.LabelFrame(main_paned, text=' 对话（可拖拽文件到此处上传，或点击「选择文件」） ', padding=4)
        main_paned.add(chat_frame, weight=1)

        self.uploaded_files = []

        # 输入区（先 pack 到底部，确保始终可见）
        self.input_frame = ttk.Frame(chat_frame, padding=(4, 8))
        self.input_frame.pack(fill=tk.X, side=tk.BOTTOM)

        # 已上传文件显示区（输入框上方）
        self.upload_display_frame = ttk.Frame(self.input_frame)
        self.upload_display_frame.pack(fill=tk.X, pady=(0, 6))
        self.upload_label_var = tk.StringVar(value='')
        ttk.Label(self.upload_display_frame, textvariable=self.upload_label_var).pack(side=tk.LEFT)
        ttk.Button(self.upload_display_frame, text='清除', command=self._clear_uploaded_files).pack(side=tk.LEFT, padx=(8, 0))

        self.input_var = tk.StringVar()
        self.input_entry = ttk.Entry(
            self.input_frame,
            textvariable=self.input_var,
            font=('Microsoft YaHei UI', 11)
        )
        self.input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 8), ipady=4)
        self.input_entry.insert(0, '在此输入问题，按 Enter 发送')
        self.input_entry.bind('<FocusIn>', lambda e: self._clear_placeholder())
        self.input_entry.bind('<Return>', lambda e: self._on_send())

        ttk.Button(self.input_frame, text='选择文件', command=self._on_choose_file).pack(side=tk.LEFT, padx=(0, 8), ipady=4)
        self.interactive_btn = ttk.Button(self.input_frame, text='交互模式', command=self._on_interactive_mode)
        self.interactive_btn.pack(side=tk.RIGHT, padx=(0, 0), ipadx=12, ipady=4)
        self.send_btn = ttk.Button(self.input_frame, text='发送', command=self._on_send)
        self.send_btn.pack(side=tk.RIGHT, ipadx=12, ipady=4)

        # 对话内容展示区（填充剩余空间）
        self.chat_text = scrolledtext.ScrolledText(
            chat_frame,
            wrap=tk.WORD,
            font=('Microsoft YaHei UI', 11),
            state=tk.DISABLED,
            bg='#fafafa',
            padx=8,
            pady=8
        )
        self.chat_text.pack(fill=tk.BOTH, expand=True)

        self.status_var = tk.StringVar(value='就绪')
        ttk.Label(self.root, textvariable=self.status_var, font=('Microsoft YaHei UI', 9)).pack(
            side=tk.BOTTOM, pady=2
        )
        _apply_ui_theme(self.root)
        self.root.after(100, lambda: _enable_drag_drop(chat_frame, self._on_files_dropped))

    def _on_interactive_mode(self):
        """点击交互模式：若思维链流程图有内容则打开新窗口。"""
        has_flow = len(self.flow_steps) > 0 or (self.flow_spec and self.flow_spec.get('nodes'))
        if not has_flow:
            messagebox.showinfo('交互模式', '思维链流程图为空，请先进行一次对话获取思维链后再使用。')
            return
        self._open_interactive_window(shen=self.shen)

    def _on_clear_all(self):
        """清空会话窗口内容和思维链流程图。"""
        self.flow_steps = []
        self.flow_spec = None
        self._flowchart_loading = False
        self.messages = []
        self.chat_text.configure(state=tk.NORMAL)
        self.chat_text.delete(1.0, tk.END)
        self.chat_text.configure(state=tk.DISABLED)
        self.interactive_btn.configure(state=tk.NORMAL)
        self._redraw_flowchart()

    def _on_save_memory(self):
        """保存到外置记忆（AI 外置记忆循环结构）。"""
        ok = _save_external_memory(self.flow_steps, self.flow_spec, self.messages)
        if ok:
            messagebox.showinfo('外置记忆', '思路已保存至外置存储器。\n路径：%s' % _EXTERNAL_MEMORY_FILE, parent=self.root)
        else:
            messagebox.showerror('外置记忆', '保存失败，请检查目录权限。', parent=self.root)

    def _on_load_memory(self):
        """加载外置记忆：弹窗输入检索描述，用 AI 匹配数据库中的流程图后加载。"""
        top = tk.Toplevel(self.root)
        top.title('加载外置记忆 - 检索')
        top.geometry('480x220')
        top.transient(self.root)
        f = ttk.Frame(top, padding=12)
        f.pack(fill=tk.BOTH, expand=True)
        ttk.Label(f, text='输入检索描述（如「关于数学证明的流程」「上次讨论的算法思路」）：').pack(anchor=tk.W)
        qvar = tk.StringVar()
        entry = ttk.Entry(f, textvariable=qvar, width=55)
        entry.pack(fill=tk.X, pady=(6, 12))
        entry.focus_set()
        status_var = tk.StringVar(value='')

        def do_load_db():
            q = qvar.get().strip()
            if not q:
                messagebox.showwarning('加载', '请输入检索描述', parent=top)
                return
            status_var.set('正在用 AI 匹配并加载…')
            top.update()

            def on_loaded(flow_steps, flow_spec):
                try:
                    top.destroy()
                except tk.TclError:
                    pass
                self.flow_steps = flow_steps
                self.flow_spec = flow_spec
                self.messages = []
                self._redraw_flowchart()
                self.chat_text.configure(state=tk.NORMAL)
                self.chat_text.delete(1.0, tk.END)
                self.chat_text.configure(state=tk.DISABLED)
                messagebox.showinfo('外置记忆', '已从数据库加载流程图。', parent=self.root)

            _load_from_database_by_query(
                self.root, q,
                self.mode_var.get(),
                self.ollama_model_var.get().strip() if self.mode_var.get() == 'ollama' else None,
                on_loaded
            )

        def do_load_file():
            loaded = _load_external_memory()
            try:
                top.destroy()
            except tk.TclError:
                pass
            if not loaded:
                messagebox.showinfo('外置记忆', '无已保存的 JSON 文件，或文件损坏。', parent=self.root)
                return
            flow_steps, flow_spec, messages = loaded
            self.flow_steps = flow_steps
            self.flow_spec = flow_spec
            self.messages = messages
            self._redraw_flowchart()
            self.chat_text.configure(state=tk.NORMAL)
            self.chat_text.delete(1.0, tk.END)
            for m in messages:
                role, content = m.get('role', ''), m.get('content', '')
                if role == 'user':
                    self.chat_text.insert(tk.END, '你：\n', 'user_tag')
                    self.chat_text.insert(tk.END, (content or '').strip() + '\n\n', 'user_msg')
                elif role == 'assistant':
                    self.chat_text.insert(tk.END, '助理：\n', 'assistant_tag')
                    self.chat_text.insert(tk.END, (content or '').strip() + '\n\n', 'assistant_msg')
            self.chat_text.configure(state=tk.DISABLED)
            self.chat_text.tag_configure('user_tag', foreground='#2d7dff')
            self.chat_text.tag_configure('assistant_tag', foreground='#0d6b0d')
            messagebox.showinfo('外置记忆', '已从 JSON 文件加载。', parent=self.root)

        btn_f = ttk.Frame(f)
        btn_f.pack(fill=tk.X)
        ttk.Button(btn_f, text='检索并加载（从数据库）', command=do_load_db).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(btn_f, text='从 JSON 文件加载', command=do_load_file).pack(side=tk.LEFT)
        ttk.Label(f, textvariable=status_var).pack(anchor=tk.W, pady=(8, 0))
        top.bind('<Return>', lambda e: do_load_db())

    def _open_interactive_window(self, shen, flow_steps=None, flow_spec=None, num_bright=None):
        """
        打开一个与主窗口布局相同的新窗口。
        shen: 记录是第几个打开的新窗口，初值为1，每点击一次继续交互加一。
        flow_steps/flow_spec: 可为 None（从主窗口复制）或传入已扩展的数据。
        num_bright: 变亮节点数量，默认 2；继续交互不输入时下一窗口为 +1。
        """
        if flow_steps is None:
            flow_steps = list(self.flow_steps)
        else:
            flow_steps = list(flow_steps)
        if flow_spec is None:
            flow_spec = (self.flow_spec.copy() if self.flow_spec else None) or {}
        else:
            flow_spec = dict(flow_spec) if flow_spec else {}
        if flow_spec.get('nodes'):
            flow_spec['nodes'] = list(flow_spec['nodes'])
            flow_spec['edges'] = list(flow_spec.get('edges', []))
        num_bright = 2 if num_bright is None else max(1, min(num_bright, 999))
        num_bright_ref = [num_bright]  # 可变引用，供 亮少/裁去 修改

        top = tk.Toplevel(self.root)
        top.title('DeepSeek 思维链对话 - 交互模式')
        def on_window_close():
            self.shen = 1
            try:
                top.destroy()
            except tk.TclError:
                pass
        top.protocol('WM_DELETE_WINDOW', on_window_close)
        try:
            geom = self.root.geometry()
            top.geometry(geom)
        except tk.TclError:
            top.geometry('1400x900')
        top.minsize(1000, 700)

        # 交互模式使用主窗口的模型选择，不再重复显示选项
        main_paned = ttk.PanedWindow(top, orient=tk.VERTICAL)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=8, pady=4)

        flowchart_frame = ttk.LabelFrame(main_paned, text=' 思维链流程图（点击后滚动鼠标缩放） ', padding=4)
        main_paned.add(flowchart_frame, weight=0)
        flow_btn_row = ttk.Frame(flowchart_frame)
        flow_btn_row.pack(fill=tk.X, pady=(0, 4))

        def save_interactive_memory():
            ok = _save_external_memory(flow_steps, flow_spec, self.messages)
            if ok:
                messagebox.showinfo('外置记忆', '思路已保存至外置存储器。', parent=top)
            else:
                messagebox.showerror('外置记忆', '保存失败。', parent=top)

        ttk.Button(flow_btn_row, text='保存外置记忆', command=save_interactive_memory).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(flow_btn_row, text='检索标签格式', command=lambda: _show_label_text_config_window(top)).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Button(flow_btn_row, text='提取', command=lambda: _show_extract_window(top, flow_steps, flow_spec)).pack(side=tk.RIGHT)
        fansi_btn = ttk.Button(flow_btn_row, text='反思')
        fansi_btn.pack(side=tk.RIGHT, padx=(0, 8))
        caiqu_btn = ttk.Button(flow_btn_row, text='裁去')
        caiqu_btn.pack(side=tk.RIGHT, padx=(0, 8))
        liangshao_btn = ttk.Button(flow_btn_row, text='亮少')
        liangshao_btn.pack(side=tk.RIGHT, padx=(0, 8))
        canvas_container = ttk.Frame(flowchart_frame)
        canvas_container.pack(fill=tk.BOTH, expand=True)
        flow_canvas = tk.Canvas(canvas_container, bg='#f5f5f5', highlightthickness=1, highlightbackground='#ccc')
        h_scroll = ttk.Scrollbar(canvas_container, orient=tk.HORIZONTAL, command=flow_canvas.xview)
        v_scroll = ttk.Scrollbar(canvas_container, orient=tk.VERTICAL, command=flow_canvas.yview)
        flow_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
        flow_canvas.configure(xscrollcommand=h_scroll.set, yscrollcommand=v_scroll.set)
        flow_zoom = [1.0]
        flow_canvas.bind('<MouseWheel>', lambda e: _interactive_zoom(flow_canvas, flow_zoom, flow_steps, flow_spec, num_bright_ref[0], e))
        flow_canvas.bind('<Button-4>', lambda e: _interactive_zoom(flow_canvas, flow_zoom, flow_steps, flow_spec, num_bright_ref[0], e, 120))
        flow_canvas.bind('<Button-5>', lambda e: _interactive_zoom(flow_canvas, flow_zoom, flow_steps, flow_spec, num_bright_ref[0], e, -120))
        flow_canvas.bind('<Enter>', lambda e: flow_canvas.focus_set())
        flowchart_frame.bind('<Configure>', lambda e: _draw_flowchart(flow_canvas, flow_steps, flow_spec, flow_zoom[0], num_bright_ref[0]))
        flow_canvas.bind('<Double-Button-1>', lambda e: _interactive_node_click(e, flow_canvas, flow_steps, flow_spec, top))
        _draw_flowchart(flow_canvas, flow_steps, flow_spec, flow_zoom[0], num_bright_ref[0])

        chat_frame = ttk.LabelFrame(main_paned, text=' 对话（可拖拽文件上传或点击「选择文件」） ', padding=4)
        main_paned.add(chat_frame, weight=1)
        input_frame = ttk.Frame(chat_frame, padding=(4, 8))
        input_frame.pack(fill=tk.X, side=tk.BOTTOM)
        uploaded_files_inter = []

        upload_display_inter = ttk.Frame(input_frame)
        upload_display_inter.pack(fill=tk.X, pady=(0, 6))
        upload_label_var_inter = tk.StringVar(value='')

        def clear_uploaded_inter():
            uploaded_files_inter.clear()
            upload_label_var_inter.set('')

        def add_uploaded_inter(paths):
            for p in paths:
                p = os.path.normpath(os.path.abspath(str(p).strip()))
                if not os.path.isfile(p):
                    continue
                ext = os.path.splitext(p)[1].lower()
                if ext not in _UPLOAD_ALLOWED_EXT:
                    messagebox.showwarning('上传', f'仅支持 .txt、.docx 及图片格式，已跳过: {os.path.basename(p)}', parent=top)
                    continue
                ok, result, err = _read_uploaded_file(p)
                if not ok:
                    messagebox.showerror('上传', f'读取失败 {os.path.basename(p)}: {err}', parent=top)
                    continue
                name = os.path.basename(p)
                if any(u['path'] == p for u in uploaded_files_inter):
                    continue
                if isinstance(result, dict) and result.get('type') == 'image':
                    uploaded_files_inter.append({'path': p, 'name': name, 'type': 'image', 'content': result['base64'], 'mime': result['mime']})
                else:
                    uploaded_files_inter.append({'path': p, 'name': name, 'type': 'text', 'content': result})
            upload_label_var_inter.set('已上传: ' + ', '.join(u['name'] for u in uploaded_files_inter) if uploaded_files_inter else '')

        ttk.Label(upload_display_inter, textvariable=upload_label_var_inter).pack(side=tk.LEFT)
        ttk.Button(upload_display_inter, text='清除', command=clear_uploaded_inter).pack(side=tk.LEFT, padx=(8, 0))

        def on_choose_file_inter():
            raw = filedialog.askopenfilenames(
                title='选择文件（支持 .txt/.docx/.png/.jpg 等）',
                filetypes=[
                    ('图片 (*.png *.jpg *.jpeg 等)', '*.png *.jpg *.jpeg *.webp *.bmp *.gif'),
                    ('PNG 图片', '*.png'),
                    ('JPEG 图片', '*.jpg'),
                    ('文本/Word', '*.txt *.docx *.doc'),
                    ('全部格式', '*.*')
                ]
            )
            paths = _normalize_file_dialog_paths(raw)
            if paths:
                add_uploaded_inter(paths)

        input_var = tk.StringVar()
        input_entry = ttk.Entry(input_frame, textvariable=input_var, font=('Microsoft YaHei UI', 11))
        input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 8), ipady=4)
        input_entry.insert(0, '在此输入问题，按 Enter 发送')
        input_entry.bind('<FocusIn>', lambda e: _clear_placeholder_interactive(input_entry))
        ttk.Button(input_frame, text='选择文件', command=on_choose_file_inter).pack(side=tk.LEFT, padx=(0, 8), ipady=4)
        continue_btn = ttk.Button(input_frame, text='继续交互')
        continue_btn.pack(side=tk.RIGHT, padx=(0, 0), ipadx=12, ipady=4)
        send_btn = ttk.Button(input_frame, text='发送')
        send_btn.pack(side=tk.RIGHT, ipadx=12, ipady=4)

        def update_continue_btn_state():
            """流程图全亮时禁用继续交互按钮。"""
            if _is_flowchart_fully_bright(flow_steps, flow_spec, num_bright_ref[0]):
                continue_btn.configure(state=tk.DISABLED)
            else:
                continue_btn.configure(state=tk.NORMAL)

        chat_text = scrolledtext.ScrolledText(chat_frame, wrap=tk.WORD, font=('Microsoft YaHei UI', 11),
                                              state=tk.NORMAL, bg='#fafafa', padx=8, pady=8)
        chat_text.pack(fill=tk.BOTH, expand=True)
        chat_text.tag_configure('user_tag', foreground='#2d7dff')
        chat_text.tag_configure('assistant_tag', foreground='#0d6b0d')
        top.after(100, lambda: _enable_drag_drop(chat_frame, add_uploaded_inter))

        status_var = tk.StringVar(value='就绪')
        ttk.Label(top, textvariable=status_var, font=('Microsoft YaHei UI', 9)).pack(side=tk.BOTTOM, pady=2)

        def append_chat(role, content, name=None):
            chat_text.configure(state=tk.NORMAL)
            if role == 'user':
                chat_text.insert(tk.END, '你：\n', 'user_tag')
                chat_text.insert(tk.END, content.strip() + '\n\n', 'user_msg')
            else:
                nm = name or ('Ollama' if self.mode_var.get() == 'ollama' else _get_cloud_assistant_name(self.mode_var.get()))
                chat_text.insert(tk.END, f'{nm}：\n', 'assistant_tag')
                chat_text.insert(tk.END, content.strip() + '\n\n', 'assistant_msg')
            chat_text.configure(state=tk.DISABLED)
            chat_text.see(tk.END)

        def redraw_flow():
            _draw_flowchart(flow_canvas, flow_steps, flow_spec, flow_zoom[0], num_bright_ref[0])

        fansi_btn.configure(command=lambda: _on_fansi(flow_steps, flow_spec, num_bright_ref, chat_text, flow_canvas, flow_zoom, top, self, status_var, send_btn, append_chat, update_continue_btn_state))
        caiqu_btn.configure(command=lambda: _on_caiqu(flow_canvas, flow_steps, flow_spec, flow_zoom, num_bright_ref, top, redraw_flow, update_continue_btn_state))
        liangshao_btn.configure(command=lambda: _on_liangshao(flow_canvas, flow_steps, flow_spec, flow_zoom, num_bright_ref, redraw_flow, update_continue_btn_state))

        def do_auto_send():
            bright_text = _get_bright_node_texts(flow_steps, flow_spec, num_bright_ref[0])
            if not bright_text.strip():
                return
            status_var.set('正在请求…')
            send_btn.configure(state=tk.DISABLED)
            mode = self.mode_var.get()
            ollama_model = self.ollama_model_var.get().strip() if mode == 'ollama' else None
            msgs = [{'role': 'user', 'content': '请仅根据以下思维链步骤给出你的回答或结论，不要重复步骤内容：\n\n' + bright_text}]
            def run():
                try:
                    if mode == 'ollama':
                        content, _ = call_ollama_api(msgs, ollama_model, use_think=False)
                        asst = ollama_model or 'Ollama'
                    else:
                        content, _ = _call_cloud_api(msgs, mode)
                        asst = _get_cloud_assistant_name(mode)
                    top.after(0, lambda: append_chat('assistant', content or '(无回复)', asst))
                except Exception as e:
                    top.after(0, lambda: messagebox.showerror('请求失败', str(e)))
                finally:
                    if _is_flowchart_fully_bright(flow_steps, flow_spec, num_bright_ref[0]):
                        mdl = ollama_model if mode == 'ollama' else _get_cloud_assistant_name(mode)
                        _om = ollama_model if mode == 'ollama' else None
                        threading.Thread(target=lambda: _save_to_database(flow_steps, flow_spec, mode, mdl, ollama_model=_om), daemon=True).start()
                    top.after(0, lambda: (send_btn.configure(state=tk.NORMAL), status_var.set('就绪'), update_continue_btn_state()))
            threading.Thread(target=run, daemon=True).start()

        def on_send():
            text = input_var.get().strip()
            if text == '在此输入问题，按 Enter 发送':
                text = ''
            text_parts = []
            image_items = []
            for u in uploaded_files_inter:
                if u.get('type') == 'image':
                    image_items.append(u)
                    text_parts.append(f'\n\n【图片 {u["name"]}】')
                else:
                    text_parts.append(f'\n\n【附件 {u["name"]}】\n{u["content"][:50000]}')
            full_content = (text + ''.join(text_parts)).strip()
            if not full_content and not image_items:
                return
            if not full_content and image_items:
                full_content = '请分析以下图片内容。'
            elif not text and text_parts and not image_items:
                full_content = '请分析以下附件内容。' + ''.join(text_parts)
            input_var.set('')
            clear_uploaded_inter()
            disp = full_content[:500] + ('…' if len(full_content) > 500 else '')
            if image_items:
                disp += ' [含 %d 张图片]' % len(image_items)
            append_chat('user', disp)
            # 先裁剪流程图：只保留当前高亮的结点，移除变暗的结点
            nb = num_bright_ref[0]
            if nb < len(flow_steps):
                flow_steps[:] = flow_steps[:nb]
            if flow_spec and flow_spec.get('nodes') and nb < len(flow_spec['nodes']):
                flow_spec['nodes'] = flow_spec['nodes'][:nb]
                edges = flow_spec.get('edges', [])
                keep_ids = {n.get('id', i + 1) for i, n in enumerate(flow_spec['nodes'])}
                flow_spec['edges'] = [e for e in edges if e.get('from') in keep_ids and e.get('to') in keep_ids]
            send_btn.configure(state=tk.DISABLED)
            continue_btn.configure(state=tk.DISABLED)
            mode = self.mode_var.get()
            ollama_model = self.ollama_model_var.get().strip() if mode == 'ollama' else None
            bright_text = _get_bright_node_texts(flow_steps, flow_spec, num_bright_ref[0])
            qa_prompt = f'''请根据以下思维链步骤和用户补充，直接给出你的回答。不要重复步骤内容，不要输出流程图或JSON，只给出面向用户的回答。

当前思维链：
{bright_text}

用户补充：{full_content}'''
            if not image_items:
                qa_msgs = [{'role': 'user', 'content': qa_prompt}]
            else:
                content_parts = [{'type': 'text', 'text': qa_prompt}]
                for img in image_items:
                    content_parts.append({'type': 'image_url', 'image_url': {'url': 'data:%s;base64,%s' % (img['mime'], img['content'])}})
                qa_msgs = [{'role': 'user', 'content': content_parts, 'images': [img['content'] for img in image_items]}]

            def run():
                def enable_buttons():
                    top.after(0, lambda: (send_btn.configure(state=tk.NORMAL), status_var.set('就绪'), update_continue_btn_state()))

                try:
                    # 第一次调用：正常对话，用户问什么答什么
                    top.after(0, lambda: status_var.set('正在请求…'))
                    if mode == 'ollama':
                        answer_content, _ = call_ollama_api(qa_msgs, ollama_model, use_think=False)
                        asst = ollama_model or 'Ollama'
                    else:
                        answer_content, _ = _call_cloud_api(qa_msgs, mode)
                        asst = _get_cloud_assistant_name(mode)
                    answer = (answer_content or '').strip()
                    top.after(0, lambda: append_chat('assistant', answer or '(无回复)', asst))

                    # 第二次调用：后台思维链续写与流程图更新，完成后才允许继续交互
                    top.after(0, lambda: status_var.set('正在更新思维链流程图…'))
                    fc_prompt = f'''基于以下思维链步骤和用户补充，请继续推导（续写思维链）。输出必须是纯流程图JSON，不要其他文字。

【重要】续写内容必须与已有思维链不重复。专业流程图规则：rect=流程步骤(长方形)，diamond=判断分支(菱形)，rounded=开始/结束(圆角)。

在回答末尾，请严格按以下格式输出续写的专业流程图JSON：
---思维链流程图JSON---
{{"nodes":[{{"id":1,"type":"rect","text":"新步骤"}}],"edges":[{{"from":1,"to":2,"label":""}}]}}
---END---

当前思维链：
{bright_text}

用户补充：{full_content}'''
                    if not image_items:
                        fc_msgs = [{'role': 'user', 'content': fc_prompt}]
                    else:
                        fc_parts = [{'type': 'text', 'text': fc_prompt}]
                        for img in image_items:
                            fc_parts.append({'type': 'image_url', 'image_url': {'url': 'data:%s;base64,%s' % (img['mime'], img['content'])}})
                        fc_msgs = [{'role': 'user', 'content': fc_parts, 'images': [img['content'] for img in image_items]}]
                    if mode == 'ollama':
                        fc_content, fc_reasoning = call_ollama_api(fc_msgs, ollama_model, use_think=False)
                    else:
                        fc_content, fc_reasoning = _call_cloud_api(fc_msgs, mode)
                    new_spec = _extract_flowchart_json_from_content(fc_content) or _extract_flowchart_json_from_content(fc_reasoning)
                    if not new_spec and fc_reasoning and len(fc_reasoning.strip()) > 50:
                        new_spec = _generate_flowchart_spec(fc_reasoning, mode, ollama_model)
                    if new_spec and new_spec.get('nodes'):
                            nb = num_bright_ref[0]
                            bright_texts = [str(n.get('text', '')) for n in flow_spec['nodes'][:nb]] if (flow_spec and flow_spec.get('nodes')) else list(flow_steps[:nb])
                            new_spec['nodes'] = _filter_redundant_nodes(new_spec['nodes'], bright_texts)
                    if new_spec and new_spec.get('nodes'):
                            # 若无 flow_spec 节点则从 flow_steps 构造
                            if not flow_spec.get('nodes') and flow_steps:
                                flow_spec['nodes'] = [{'id': i + 1, 'type': 'rect', 'text': s} for i, s in enumerate(flow_steps)]
                                flow_spec['edges'] = [{'from': i, 'to': i + 1, 'label': ''} for i in range(1, len(flow_steps))]
                            nodes = flow_spec['nodes']
                            last_id = max(n.get('id', i + 1) for i, n in enumerate(nodes))
                            base_id = last_id + 1
                            new_nodes = new_spec['nodes']
                            id_map = {n.get('id', i + 1): base_id + i for i, n in enumerate(new_nodes)}
                            flow_spec.setdefault('edges', []).append({'from': last_id, 'to': base_id, 'label': ''})
                            for i, n in enumerate(new_nodes):
                                flow_spec['nodes'].append({
                                    'id': base_id + i,
                                    'type': n.get('type', 'rect'),
                                    'text': n.get('text', '')
                                })
                            for e in new_spec.get('edges', []):
                                f, t = e.get('from'), e.get('to')
                                if f in id_map and t in id_map:
                                    flow_spec['edges'].append({'from': id_map[f], 'to': id_map[t], 'label': e.get('label', '')})
                            for n in new_nodes:
                                flow_steps.append(str(n.get('text', '')))
                            top.after(0, redraw_flow)
                except Exception as e:
                    top.after(0, lambda: messagebox.showerror('请求失败', str(e)))
                finally:
                    top.after(0, lambda: _save_external_memory(flow_steps, flow_spec, self.messages))
                    if _is_flowchart_fully_bright(flow_steps, flow_spec, num_bright_ref[0]):
                        mdl = ollama_model if mode == 'ollama' else _get_cloud_assistant_name(mode)
                        _om = ollama_model if mode == 'ollama' else None
                        threading.Thread(target=lambda: _save_to_database(flow_steps, flow_spec, mode, mdl, ollama_model=_om), daemon=True).start()
                    enable_buttons()
            threading.Thread(target=run, daemon=True).start()

        def on_continue():
            self.shen = shen + 1
            try:
                top.destroy()
            except tk.TclError:
                pass
            nb = num_bright_ref[0]
            next_bright = min(nb + 1, len(flow_steps) or 999)
            if flow_spec and flow_spec.get('nodes'):
                next_bright = min(nb + 1, len(flow_spec['nodes']))
            self._open_interactive_window(shen=self.shen, flow_steps=flow_steps, flow_spec=flow_spec, num_bright=next_bright)

        send_btn.configure(command=on_send)
        input_entry.bind('<Return>', lambda e: on_send())
        continue_btn.configure(command=on_continue)
        chat_text.configure(state=tk.DISABLED)
        top.after(100, do_auto_send)
        top.after(150, update_continue_btn_state)

    def _on_mode_change(self):
        if self.mode_var.get() == 'ollama':
            self.ollama_frame.pack(side=tk.LEFT, padx=(8, 0))
            self._refresh_ollama_models()
        else:
            self.ollama_frame.pack_forget()

    def _refresh_ollama_models(self):
        def do_fetch():
            try:
                models = fetch_ollama_models()
                self.root.after(0, lambda: self._apply_ollama_models(models))
            except Exception as e:
                self.root.after(0, lambda: self._on_ollama_refresh_error(str(e)))
        threading.Thread(target=do_fetch, daemon=True).start()
        self.status_var.set('正在获取 Ollama 模型列表…')

    def _apply_ollama_models(self, models):
        self.status_var.set('就绪')
        self.model_combo['values'] = models
        if models and not self.ollama_model_var.get():
            # 优先选择 deepseek-r1 或包含 deepseek 的模型
            for m in models:
                if 'deepseek' in m.lower():
                    self.ollama_model_var.set(m)
                    return
            self.ollama_model_var.set(models[0])

    def _on_flowchart_zoom(self, event, delta_override=None):
        delta = delta_override if delta_override is not None else (event.delta if hasattr(event, 'delta') else 0)
        if delta == 0:
            return
        self.flow_zoom *= (1.1 if delta > 0 else 0.9)
        self.flow_zoom = max(0.5, min(2.0, self.flow_zoom))
        self._redraw_flowchart()

    def _redraw_flowchart(self):
        _draw_flowchart(self.flow_canvas, self.flow_steps, self.flow_spec, self.flow_zoom,
                       loading=self._flowchart_loading)

    def _on_flowchart_node_double_click(self, event):
        ids = self.flow_canvas.find_closest(event.x, event.y)
        if not ids:
            return
        tags = self.flow_canvas.gettags(ids[0])
        node_tag = next((t for t in tags if t.startswith('node_')), None)
        if not node_tag:
            return
        try:
            idx = int(node_tag.split('_')[1])
        except (ValueError, IndexError):
            return
        full_text = ''
        if self.flow_spec and self.flow_spec.get('nodes'):
            nodes = self.flow_spec['nodes']
            if 0 <= idx < len(nodes):
                full_text = str(nodes[idx].get('text', ''))
        elif self.flow_steps and 0 <= idx < len(self.flow_steps):
            full_text = self.flow_steps[idx]
        if not full_text:
            return
        top = tk.Toplevel(self.root)
        top.title('思维链节点 - 完整内容')
        top.geometry('500x360')
        top.minsize(400, 200)

        def on_close():
            try:
                top.destroy()
            except tk.TclError:
                pass
            try:
                self.root.unbind('<Button-1>', on_root_click)
            except (tk.TclError, AttributeError, NameError):
                pass

        def on_root_click(ev):
            try:
                if top.winfo_exists():
                    on_close()
            except tk.TclError:
                pass

        top.protocol('WM_DELETE_WINDOW', on_close)

        btn_frame = ttk.Frame(top)
        btn_frame.pack(fill=tk.X, padx=8, pady=4)
        btn = ttk.Button(btn_frame, text='关闭 (或按 Esc / 点击主窗口)', command=on_close, width=28)
        btn.pack(side=tk.RIGHT)

        txt = tk.Text(top, wrap=tk.WORD, font=('Microsoft YaHei UI', 10), padx=8, pady=8)
        txt.pack(fill=tk.BOTH, expand=True)
        txt.insert(tk.END, full_text)
        txt.configure(state=tk.DISABLED)

        top.bind('<Escape>', lambda e: on_close())
        self.root.bind('<Button-1>', on_root_click, add='+')

    def _apply_professional_flowchart(self, flow_spec):
        self.flow_spec = flow_spec
        self._redraw_flowchart()
        _save_external_memory(self.flow_steps, self.flow_spec, self.messages)

    def _on_ollama_refresh_error(self, err):
        self.status_var.set('Ollama 未就绪')
        self.model_combo['values'] = []
        messagebox.showwarning('Ollama 不可用', f'无法获取 Ollama 模型列表：\n{err}\n\n请确保 Ollama 已安装并运行。')

    def _clear_placeholder(self):
        if self.input_entry.get() == '在此输入问题，按 Enter 发送':
            self.input_entry.delete(0, tk.END)

    def _on_choose_file(self):
        raw = filedialog.askopenfilenames(
            title='选择文件（支持 .txt/.docx/.png/.jpg 等）',
            filetypes=[
                ('图片 (*.png *.jpg *.jpeg 等)', '*.png *.jpg *.jpeg *.webp *.bmp *.gif'),
                ('PNG 图片', '*.png'),
                ('JPEG 图片', '*.jpg'),
                ('文本/Word', '*.txt *.docx *.doc'),
                ('全部格式', '*.*')
            ]
        )
        paths = _normalize_file_dialog_paths(raw)
        if paths:
            self._add_uploaded_files(paths)

    def _on_files_dropped(self, paths):
        self._add_uploaded_files(paths)

    def _add_uploaded_files(self, paths):
        for p in paths:
            p = os.path.normpath(os.path.abspath(str(p).strip()))
            if not os.path.isfile(p):
                continue
            ext = os.path.splitext(p)[1].lower()
            if ext not in _UPLOAD_ALLOWED_EXT:
                messagebox.showwarning('上传', f'仅支持 .txt、.docx 及图片格式(.png/.jpg 等)，已跳过: {os.path.basename(p)}', parent=self.root)
                continue
            ok, result, err = _read_uploaded_file(p)
            if not ok:
                messagebox.showerror('上传', f'读取失败 {os.path.basename(p)}: {err}', parent=self.root)
                continue
            name = os.path.basename(p)
            if any(u['path'] == p for u in self.uploaded_files):
                continue
            if isinstance(result, dict) and result.get('type') == 'image':
                self.uploaded_files.append({
                    'path': p, 'name': name, 'type': 'image',
                    'content': result['base64'], 'mime': result['mime']
                })
            else:
                self.uploaded_files.append({'path': p, 'name': name, 'type': 'text', 'content': result})
        self._update_upload_display()

    def _clear_uploaded_files(self):
        self.uploaded_files.clear()
        self._update_upload_display()

    def _update_upload_display(self):
        try:
            if not self.uploaded_files:
                self.upload_label_var.set('')
            else:
                names = [u['name'] for u in self.uploaded_files]
                self.upload_label_var.set('已上传: ' + ', '.join(names))
            self.upload_display_frame.update_idletasks()
            self.root.update_idletasks()
        except (tk.TclError, AttributeError):
            pass

    def _append_chat(self, role, content, assistant_name=None):
        self.chat_text.configure(state=tk.NORMAL)
        if role == 'user':
            self.chat_text.insert(tk.END, '你：\n', 'user_tag')
            self.chat_text.insert(tk.END, content.strip() + '\n\n', 'user_msg')
        else:
            name = assistant_name or ('Ollama' if self.mode_var.get() == 'ollama' else _get_cloud_assistant_name(self.mode_var.get()))
            self.chat_text.insert(tk.END, f'{name}：\n', 'assistant_tag')
            self.chat_text.insert(tk.END, content.strip() + '\n\n', 'assistant_msg')
        self.chat_text.configure(state=tk.DISABLED)
        self.chat_text.see(tk.END)
        self.chat_text.tag_configure('user_tag', foreground='#2d7dff')
        self.chat_text.tag_configure('assistant_tag', foreground='#0d6b0d')

    def _on_send(self):
        text = self.input_var.get().strip()
        if text == '在此输入问题，按 Enter 发送':
            text = ''
        text_parts = []
        image_items = []
        for u in self.uploaded_files:
            if u.get('type') == 'image':
                image_items.append(u)
                text_parts.append(f'\n\n【图片 {u["name"]}】')
            else:
                text_parts.append(f'\n\n【附件 {u["name"]}】\n{u["content"][:50000]}')
        full_text = (text + ''.join(text_parts)).strip()
        if not full_text and not image_items:
            return
        if not full_text and image_items:
            full_text = '请分析以下图片内容。'
        elif not text and text_parts and not image_items:
            full_text = '请分析以下附件内容。' + ''.join(text_parts)

        self.input_var.set('')
        self._clear_uploaded_files()
        display_text = full_text[:500] + ('…' if len(full_text) > 500 else '')
        if image_items:
            display_text += ' [含 %d 张图片]' % len(image_items)
        self._append_chat('user', display_text)

        # 构建 API 消息：云端用 content 数组（OpenAI 多模态格式），Ollama 用 content + images
        if not image_items:
            user_msg = {'role': 'user', 'content': full_text}
        else:
            content_parts = [{'type': 'text', 'text': full_text}]
            for img in image_items:
                content_parts.append({
                    'type': 'image_url',
                    'image_url': {'url': 'data:%s;base64,%s' % (img['mime'], img['content'])}
                })
            user_msg = {'role': 'user', 'content': content_parts, 'images': [img['content'] for img in image_items]}
        api_messages = self.messages + [user_msg]
        self.messages.append({'role': 'user', 'content': full_text})
        self.send_btn.configure(state=tk.DISABLED)
        mode = self.mode_var.get()
        ollama_model = self.ollama_model_var.get().strip() if mode == 'ollama' else None
        self.status_var.set('正在请求…' if mode == 'ollama' else '正在请求 %s…' % _get_cloud_assistant_name(mode))

        def do_request():
            try:
                if mode == 'ollama':
                    content, reasoning = call_ollama_api(api_messages, ollama_model, use_think=False)
                    asst_name = ollama_model or 'Ollama'
                else:
                    content, reasoning = _call_cloud_api(api_messages, mode)
                    asst_name = _get_cloud_assistant_name(mode)

                self.messages.append({'role': 'assistant', 'content': content})
                self.root.after(0, lambda: self._append_chat('assistant', content or '(无回复内容)', asst_name))
                steps = fangfa1(reasoning)
                # 不先显示简单流程图，等专业流程图生成后再显示；生成期间展示加载提示
                self.root.after(0, lambda: self._show_flowchart_loading())
                if reasoning and len(reasoning.strip()) > 50:
                    def gen_flowchart():
                        try:
                            fs = _extract_flowchart_json_from_content(content) or _extract_flowchart_json_from_content(reasoning)
                            if not fs:
                                fs = _generate_flowchart_spec(reasoning, mode, ollama_model)
                            if fs and fs.get('nodes'):
                                flow_steps_from_spec = [str(n.get('text', '')) for n in fs['nodes']]
                                self.root.after(0, lambda: self._apply_flowchart_done(flow_steps_from_spec, fs))
                            else:
                                self.root.after(0, lambda: self._apply_flowchart_done(steps, None))
                        except Exception:
                            self.root.after(0, lambda: self._apply_flowchart_done(steps, None))
                    threading.Thread(target=gen_flowchart, daemon=True).start()
                else:
                    self.root.after(0, lambda: self._apply_flowchart_done(steps, None))
            except Exception as e:
                err = str(e)
                # 网络相关错误给出更清晰提示
                if any(k in err.lower() for k in ('connection', 'connect', 'network', 'timeout', '无法连接', '超时')):
                    err = f'{err}\n\n若未联网，请先连接网络后再试。'
                self.root.after(0, lambda: self._on_error(err))
            finally:
                self.root.after(0, lambda: self.send_btn.configure(state=tk.NORMAL))

        threading.Thread(target=do_request, daemon=True).start()

    def _show_flowchart_loading(self):
        """专业流程图生成中：清空流程图区域，显示加载提示，禁用交互模式。若已暂停刷新则不清空。"""
        self._flowchart_loading = True
        if not self._flowchart_refresh_paused:
            self.flow_steps = []
            self.flow_spec = None
            self.status_var.set('正在生成专业流程图…')
        else:
            self.status_var.set('正在生成专业流程图…（已暂停刷新，当前流程图保持不变）')
        self.interactive_btn.configure(state=tk.DISABLED)
        self._redraw_flowchart()

    def _apply_flowchart_done(self, flow_steps, flow_spec):
        """专业流程图就绪（或兜底为简单流程图）：更新显示，恢复交互模式。若已暂停刷新则不更新流程图内容。"""
        self._flowchart_loading = False
        if not self._flowchart_refresh_paused:
            self.flow_steps = flow_steps or []
            self.flow_spec = flow_spec
        self.status_var.set('就绪')
        self.interactive_btn.configure(state=tk.NORMAL)
        self._redraw_flowchart()
        _save_external_memory(self.flow_steps, self.flow_spec, self.messages)

    def _update_after_reply(self, content, steps, assistant_name=None, flow_spec=None):
        self.status_var.set('就绪')
        if not self._flowchart_refresh_paused:
            self.flow_steps = steps
            self.flow_spec = flow_spec
        self._redraw_flowchart()
        self._append_chat('assistant', content or '(无回复内容)', assistant_name)
        _save_external_memory(self.flow_steps, self.flow_spec, self.messages)

    def _on_pause_refresh(self):
        """暂停刷新：禁止自动更新思维链流程图。"""
        self._flowchart_refresh_paused = True
        self.pause_refresh_btn.configure(state=tk.DISABLED)
        self.resume_refresh_btn.configure(state=tk.NORMAL)
        self.status_var.set('已暂停刷新流程图')

    def _on_resume_refresh(self):
        """继续刷新：恢复正常的流程图自动更新。"""
        self._flowchart_refresh_paused = False
        self.pause_refresh_btn.configure(state=tk.NORMAL)
        self.resume_refresh_btn.configure(state=tk.DISABLED)
        self.status_var.set('就绪')

    def _do_zisikao_direct(self):
        """直接自思考：流程图内容 + 会话内容 一并发给模型，等待一次返回。"""
        flow_text = _extract_flowchart_as_text(self.flow_steps, self.flow_spec)
        if not flow_text:
            messagebox.showinfo('自思考', '思维链流程图为空。', parent=self.root)
            return
        chat_content = self.chat_text.get(1.0, tk.END).strip()
        parts = ['请按照以下思维链流程图的步骤执行，并给出最终结果。不要重复步骤内容，只输出面向用户的结论或回答。\n\n思维链流程图：\n', flow_text]
        if chat_content:
            parts.append('\n\n当前会话内容（请结合上述流程图一并考虑）：\n')
            parts.append(chat_content)
        prompt = ''.join(parts)
        self.send_btn.configure(state=tk.DISABLED)
        self.status_var.set('正在自思考…')

        def run():
            try:
                mode = self.mode_var.get()
                ollama_model = self.ollama_model_var.get().strip() if mode == 'ollama' else None
                msgs = [{'role': 'user', 'content': prompt}]
                if mode == 'ollama':
                    content, _ = call_ollama_api(msgs, ollama_model, use_think=False)
                    asst = ollama_model or 'Ollama'
                else:
                    content, _ = _call_cloud_api(msgs, mode)
                    asst = _get_cloud_assistant_name(mode)
                self.root.after(0, lambda: self._append_chat('assistant', content or '(无回复)', asst))
            except Exception as e:
                self.root.after(0, lambda: self._on_error(str(e)))
            finally:
                self.root.after(0, lambda: (self.send_btn.configure(state=tk.NORMAL), self.status_var.set('就绪')))

        threading.Thread(target=run, daemon=True).start()

    def _do_zisikao_loop(self):
        """循环自思考：按节点顺序，每一步的输入 = 上一步输出 + 当前节点内容+功能，最后结果放入会话。"""
        nodes = _get_flowchart_nodes_ordered(self.flow_steps, self.flow_spec)
        if not nodes:
            messagebox.showinfo('自思考', '思维链流程图为空。', parent=self.root)
            return
        chat_content = self.chat_text.get(1.0, tk.END).strip()
        accum = chat_content if chat_content else '（无初始内容）'
        self.send_btn.configure(state=tk.DISABLED)
        self.status_var.set('正在循环自思考…')

        def run():
            try:
                mode = self.mode_var.get()
                ollama_model = self.ollama_model_var.get().strip() if mode == 'ollama' else None
                asst = ollama_model or 'Ollama' if mode == 'ollama' else _get_cloud_assistant_name(mode)
                for i, (node_text, node_func) in enumerate(nodes):
                    self.root.after(0, lambda n=i+1, t=len(nodes): self.status_var.set('正在自思考 节点 %d/%d…' % (n, t)))
                    prompt = f'''当前输入/上一步输出：
{accum}

当前节点（{node_func}）：
{node_text}

请根据当前节点的功能和内容，基于上述输入进行处理，只输出处理后的结果，不要其他解释。'''
                    msgs = [{'role': 'user', 'content': prompt}]
                    if mode == 'ollama':
                        content, _ = call_ollama_api(msgs, ollama_model, use_think=False)
                    else:
                        content, _ = _call_cloud_api(msgs, mode)
                    accum = (content or '').strip() or accum
                self.root.after(0, lambda: self._append_chat('assistant', accum, asst))
            except Exception as e:
                self.root.after(0, lambda: self._on_error(str(e)))
            finally:
                self.root.after(0, lambda: (self.send_btn.configure(state=tk.NORMAL), self.status_var.set('就绪')))

        threading.Thread(target=run, daemon=True).start()

    def _on_error(self, err_msg):
        self.send_btn.configure(state=tk.NORMAL)
        self.status_var.set('请求失败')
        try:
            safe_msg = str(err_msg).encode('utf-8', errors='replace').decode('utf-8')
        except Exception:
            safe_msg = '未知错误'
        suffix = '若为 Ollama 模式，请确保 Ollama 已启动。' if self.mode_var.get() == 'ollama' else '若为云端模式，请检查网络与 API Key（双击对应模型可输入 Key）。'
        messagebox.showerror('请求失败', f'调用失败：\n{safe_msg}\n\n{suffix}')

    def run(self):
        self.root.mainloop()


if __name__ == '__main__':
    app = DeepSeekChatApp()
    app.run()
